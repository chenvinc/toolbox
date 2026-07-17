"""JSON→Word 试卷生成端到端测试（数据管线闭环）。

用真实样本 docs/专项智能练习（数量关系）_2026-07-17T06-01-08-897.json 作为输入，
经服务层（JsonToWordServiceImpl + DocxExamWriterAdapter）跑完整闭环：
  解析 → 并发预下载图片 → 生成题本 → 生成解析 → 结果。

为离线/确定性，图片下载统一 mock 为本地 PNG 字节；场景 C 再让其中部分 URL 失败。

覆盖三个场景：
  A. 默认排版参数生成
  B. 自定义行间距 + 关闭首行缩进生成
  C. 模拟部分图片链接失效
并验证：输出文件名（{pageTitle}_题本/_解析）、题数、步骤进度事件、失败图片列表、
并发预下载机制、以及输出目录无权限时的 OutputWriteError。
"""
import base64
import json
import os
import tempfile
import unittest
from pathlib import Path
from typing import Dict, List, Optional
from unittest.mock import patch

from shared.contracts import (
    EventType, ExamLineSpacingType, GenerateExamRequest, GenerateExamResult,
)
from shared.errors import OutputWriteError
from core.models.exam_question import ExamImage, ExamQuestion
from core.services._exam_image_fetch import prefetch_images
from core.services.json_to_word_service import JsonToWordServiceImpl
from core.adapters.docx_exam_writer import DocxExamWriterAdapter
from docx import Document

_SAMPLE = (Path(__file__).resolve().parents[3] / "生成示例" /
           "专项智能练习（数量关系）_2026-07-17T05-56-50-992.json")

# 8x6 有效 PNG（与适配器测试同源），作为 mock 下载结果，避免依赖外部网络。
_PNG_B64 = "iVBORw0KGgoAAAANSUhEUgAAAAgAAAAGCAIAAABxZ0isAAAAEklEQVR4nGP8z4AdMOEQp4sEAJiFAQtsMIKbAAAAAElFTkSuQmCC"
_PNG_BYTES = base64.b64decode(_PNG_B64)


class CollectingEmitter:
    """实现 EventEmitter 端口，收集所有事件供断言。"""

    def __init__(self):
        self.events = []
        self._handlers = []

    def emit(self, event):
        self.events.append(event)
        for h in self._handlers:
            h(event)

    def on_event(self, handler):
        self._handlers.append(handler)


def _unique_image_urls(path: str) -> List[str]:
    """从样本 JSON 中按服务相同的顺序收集去重图片 URL。"""
    with open(path, "r", encoding="utf-8") as fh:
        data = json.load(fh)
    seen: set = set()
    urls: List[str] = []
    for q in data.get("questions", []):
        for im in (q.get("images") or []):
            s = (im.get("src") or "").strip()
            if s and s not in seen:
                seen.add(s)
                urls.append(s)
    return urls


def _all_ok(src, cache, lock=None):
    """mock：任意 URL 均下载成功，返回本地 PNG 字节。"""
    if src in cache:
        return cache[src]
    cache[src] = _PNG_BYTES
    return _PNG_BYTES


def _fail_subset(fail_set):
    """mock：对落入 fail_set 的 URL 返回 None（下载失败），其余成功。"""

    def _fetch(src, cache, lock=None):
        if src in cache:
            return cache[src]
        data = None if src in fail_set else _PNG_BYTES
        cache[src] = data
        return data

    return _fetch


class JsonExamE2ETests(unittest.TestCase):
    def setUp(self):
        if not _SAMPLE.exists():
            self.skipTest(f"样本文件缺失：{_SAMPLE}")
        self._dir = tempfile.mkdtemp()
        self._out = os.path.join(self._dir, "out")
        os.makedirs(self._out)
        self.emitter = CollectingEmitter()
        self.writer = DocxExamWriterAdapter()
        self.svc = JsonToWordServiceImpl(self.writer, self.emitter)
        self._urls = _unique_image_urls(str(_SAMPLE))
        # 让索引为偶数的 URL 失败（约半数），用于场景 C。
        self._fail_set = {u for i, u in enumerate(self._urls) if i % 2 == 0}
        # 题数从样本直接读取，避免对具体样本硬编码。
        with open(str(_SAMPLE), "r", encoding="utf-8") as fh:
            self._expected_count = len(json.load(fh).get("questions", []))

    def _default_req(self) -> GenerateExamRequest:
        return GenerateExamRequest(input_path=str(_SAMPLE), output_dir=self._out)

    # ── 场景 A：默认排版参数 ──
    def test_scenario_a_default_params(self):
        req = self._default_req()
        with patch("core.services._exam_image_fetch._fetch_image_bytes", _all_ok):
            result = self.svc.generate(req)

        self.assertIsInstance(result, GenerateExamResult)
        self.assertEqual(result.question_count, self._expected_count)
        self.assertEqual(result.failed_images, [])  # 全部下载成功
        self.assertTrue(os.path.exists(result.question_book_path))
        self.assertTrue(os.path.exists(result.analysis_path))

        # 输出文件名：{pageTitle}_题本.docx / {pageTitle}_解析.docx
        self.assertTrue(
            result.question_book_path.endswith("专项智能练习（数量关系）_题本.docx")
        )
        self.assertTrue(
            result.analysis_path.endswith("专项智能练习（数量关系）_解析.docx")
        )

        # 步骤进度事件齐备：解析中 → 下载图片 → 生成题本 → 生成解析 → 完成
        msgs = [e.message for e in self.emitter.events
                if e.type == EventType.EXAM_PROGRESS]
        self.assertTrue(any(m.startswith("解析中") for m in msgs))
        self.assertTrue(any(m.startswith("下载图片") for m in msgs))
        self.assertIn("生成题本...", msgs)
        self.assertIn("生成解析...", msgs)
        self.assertTrue(
            any(e.type == EventType.EXAM_COMPLETED for e in self.emitter.events)
        )

    # ── 场景 B：自定义行间距 1.8 + 关闭首行缩进 ──
    def test_scenario_b_custom_spacing_no_indent(self):
        req = GenerateExamRequest(
            input_path=str(_SAMPLE), output_dir=self._out,
            line_spacing_type=ExamLineSpacingType.CUSTOM,
            line_spacing_value=1.8,
            first_line_indent=False,
        )
        with patch("core.services._exam_image_fetch._fetch_image_bytes", _all_ok):
            result = self.svc.generate(req)

        self.assertTrue(os.path.exists(result.question_book_path))
        doc = Document(result.question_book_path)
        # 首个非空段落为题本第一题题干，应套用 1.8 倍行距且无首行缩进
        first = next(p for p in doc.paragraphs if p.text.strip())
        self.assertAlmostEqual(
            float(first.paragraph_format.line_spacing), 1.8, places=3
        )
        indent = first.paragraph_format.first_line_indent
        self.assertIsNotNone(indent)
        self.assertAlmostEqual(float(indent), 0.0, places=3)

    # ── 场景 C：模拟部分图片链接失效 ──
    def test_scenario_c_partial_image_failure(self):
        req = self._default_req()
        with patch("core.services._exam_image_fetch._fetch_image_bytes",
                   _fail_subset(self._fail_set)):
            result = self.svc.generate(req)

        # 生成仍成功（占位框替代），题数与文件完整
        self.assertEqual(result.question_count, self._expected_count)
        self.assertTrue(os.path.exists(result.question_book_path))
        self.assertTrue(os.path.exists(result.analysis_path))

        # 失败图片列表与预期的失败子集一致
        self.assertEqual(set(result.failed_images), self._fail_set)
        self.assertTrue(len(result.failed_images) > 0)

        # 进度栏出现部分失败警告（下载失败与「下载成功但字节不可识别」均计入）
        msgs = [e.message for e in self.emitter.events
                if e.type == EventType.EXAM_PROGRESS]
        self.assertTrue(any("部分图片未能加载" in m for m in msgs))

        # 题本文档出现灰色占位框文字
        doc = Document(result.question_book_path)
        self.assertTrue(
            any("[图片加载失败:" in p.text for p in doc.paragraphs)
        )

    # ── 并发预下载机制（独立单元验证） ──
    def test_prefetch_collects_failures_offline(self):
        urls = ["https://x/1.png", "https://x/2.png", "https://x/3.png"]
        cache, failed = prefetch_images(urls, on_progress=lambda *_: None)
        # 离线环境下这些 URL 必失败
        self.assertEqual(set(failed), set(urls))
        self.assertEqual(set(cache.keys()), set(urls))
        for v in cache.values():
            self.assertIsNone(v)

    # ── 输出目录无写入权限：抛 OutputWriteError ──
    def test_permission_error_raises_output_write_error(self):
        req = self._default_req()
        with patch("core.services._exam_image_fetch._fetch_image_bytes", _all_ok):
            with patch("core.adapters.docx_exam_writer.os.makedirs",
                       side_effect=PermissionError("denied")):
                with self.assertRaises(OutputWriteError):
                    self.svc.generate(req)


class ImageFailureReportingTests(unittest.TestCase):
    """回归测试：下载「成功」但字节不可识别（如源站返回加密/非图片内容）的图片，
    必须如实计入 failed_images 并插入占位框，不能让 UI 误报「全部成功」。

    不依赖外部样本：直接构造题目与缓存，离线确定性。
    """

    def setUp(self):
        self._dir = tempfile.mkdtemp()
        self._out = os.path.join(self._dir, "out")
        os.makedirs(self._out)

    def _build_one(self, cache):
        req = GenerateExamRequest(
            input_path=os.path.join(self._dir, "x.json"), output_dir=self._out,
        )
        q = ExamQuestion(
            number="1.", question_type="单选题",
            stem="题干含图 [IMG1]", options={"A": "A. 1"},
            correct_answer="A", correct_rate="50%", analysis="解析 [IMG1]",
            images=[ExamImage(index=1, src="https://x/bad.png", role="stem", is_tex=False)],
        )
        return DocxExamWriterAdapter().build(
            req, [q], on_progress=lambda *_: None, image_cache=cache,
        )

    def test_unrenderable_bytes_reported_as_failed(self):
        # 缓存里是非图片字节（模拟 fb.fenbike.cn 返回的加密/非图片内容）。
        cache = {"https://x/bad.png": b"<html>not an image</html>"}
        result = self._build_one(cache)
        self.assertEqual(result.failed_images, ["https://x/bad.png"])

        book = Document(result.question_book_path)
        self.assertTrue(any("[图片加载失败:" in p.text for p in book.paragraphs))

    def test_valid_bytes_not_reported(self):
        # 合法 PNG：不应计入失败。
        cache = {"https://x/bad.png": _PNG_BYTES}
        result = self._build_one(cache)
        self.assertEqual(result.failed_images, [])


if __name__ == "__main__":
    unittest.main()
