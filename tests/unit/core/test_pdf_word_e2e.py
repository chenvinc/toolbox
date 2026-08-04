"""PdfWordConverterAdapter 端到端测试（真实 pymupdf / python-docx，临时文件）。

验证：
  - 页数与源 PDF 一致、空页统计正确；
  - 文字内容按阅读顺序保留（流式段落）；
  - 可选模板复用：清空模板原有正文后写入转换结果；
  - 行内 run 的粗体 / 红色保留；
  - 非法 PDF 抛 PdfReadError。
"""
import os
import tempfile
import unittest

import fitz
from docx import Document
from docx.dml.color import RGBColor

from core.adapters.pdf_word_converter import PdfWordConverterAdapter
from shared.contracts import ConvertPdfToWordRequest
from shared.errors import PdfReadError


def _make_pdf(path, with_style=False):
    """生成 3 页 PDF：第 2 页为空（无文字）。"""
    doc = fitz.open()
    p1 = doc.new_page(width=720, height=405)
    p1.insert_text((72, 72), "Hello PDF", fontsize=18)
    p1.insert_text((72, 102), "Second line", fontsize=18)
    doc.new_page(width=720, height=405)  # 空页
    p3 = doc.new_page(width=720, height=405)
    p3.insert_text((72, 72), "Page three", fontsize=18)
    if with_style:
        # 粗体 + 红色，用于校验 run 级格式保留
        p3.insert_text(
            (72, 102), "BoldRed", fontsize=18,
            fontname="Helvetica-Bold", color=(1, 0, 0),
        )
    doc.save(path)
    doc.close()


def _make_template(path):
    """生成带标记段落的 Word 模板（验证复用时会清空）。"""
    d = Document()
    d.add_paragraph("TEMPLATE_MARKER_SHOULD_BE_CLEARED")
    d.save(path)


def _make_cjk_pdf(path):
    """生成单页 PDF：一段会被窄栏强制换行的中文长句（多视觉行）。"""
    doc = fitz.open()
    page = doc.new_page(width=720, height=405)
    rect = fitz.Rect(72, 72, 300, 320)  # 窄栏 → 自动换行成多行
    text = (
        "中文段落重排测试：这是一段较长的说明文字，"
        "转换器应当把被PDF换行拆散的视觉行，"
        "重新合并为连贯的单个Word段落。"
    )
    page.insert_textbox(rect, text, fontsize=16, fontname="china-s")
    doc.save(path)
    doc.close()
    return text


def _make_en_pdf(path):
    """生成单页 PDF：一段会被窄栏强制换行的英文长句（多视觉行）。"""
    doc = fitz.open()
    page = doc.new_page(width=720, height=405)
    rect = fitz.Rect(72, 72, 260, 320)  # 窄栏 → 自动换行成多行
    text = (
        "This is a fairly long English sentence that should be wrapped "
        "across several visual lines by the converter tool chain."
    )
    page.insert_textbox(rect, text, fontsize=16)
    doc.save(path)
    doc.close()
    return text


def _collect(docx_path):
    d = Document(docx_path)
    paras = [p.text for p in d.paragraphs]
    return paras, d


class PdfWordConverterE2ETests(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self._tmp.cleanup)
        self.pdf_path = os.path.join(self._tmp.name, "in.pdf")
        self.tpl_path = os.path.join(self._tmp.name, "tpl.docx")
        self.out_path = os.path.join(self._tmp.name, "out.docx")
        _make_pdf(self.pdf_path)
        _make_template(self.tpl_path)
        self.adapter = PdfWordConverterAdapter()

    def _convert(self, template_path=""):
        progress = []
        result = self.adapter.convert(
            ConvertPdfToWordRequest(
                pdf_path=self.pdf_path,
                template_path=template_path,
                output_path=self.out_path,
            ),
            lambda c, t: progress.append((c, t)),
        )
        return result, progress

    def test_page_count_and_empty_pages(self):
        result, progress = self._convert()
        self.assertEqual(result.page_count, 3)
        self.assertEqual(result.empty_pages, [2])
        self.assertGreater(result.paragraph_count, 0)
        self.assertGreater(result.run_count, 0)
        self.assertEqual(progress, [(1, 3), (2, 3), (3, 3)])

    def test_text_content_preserved(self):
        self._convert()
        paras, _ = _collect(self.out_path)
        joined = "\n".join(paras)
        self.assertIn("Hello PDF", joined)
        self.assertIn("Page three", joined)

    def test_template_reuse_clears_original_content(self):
        # 套用模板：模板原有段落应被清空，仅留转换结果
        self._convert(template_path=self.tpl_path)
        paras, _ = _collect(self.out_path)
        joined = "\n".join(paras)
        self.assertNotIn("TEMPLATE_MARKER_SHOULD_BE_CLEARED", joined)
        self.assertIn("Hello PDF", joined)

    def test_bold_and_color_preserved(self):
        _make_pdf(self.pdf_path, with_style=True)
        self._convert()
        _, d = _collect(self.out_path)
        found = False
        for p in d.paragraphs:
            for run in p.runs:
                if run.text == "BoldRed":
                    found = True
                    self.assertTrue(run.font.bold, "粗体应保留")
                    self.assertEqual(str(run.font.color.rgb), "FF0000", "红色应保留")
        self.assertTrue(found, "应存在 BoldRed 文本")

    def test_chinese_paragraph_reconstruction(self):
        # 中文长段落：多视觉行应被合并为单一 Word 段落，且中文间不插空格。
        cjk_path = os.path.join(self._tmp.name, "cjk.pdf")
        text = _make_cjk_pdf(cjk_path)
        self.adapter.convert(
            ConvertPdfToWordRequest(
                pdf_path=cjk_path, template_path="", output_path=self.out_path
            ),
            lambda c, t: None,
        )
        paras, _ = _collect(self.out_path)
        non_empty = [p for p in paras if p.strip()]
        self.assertEqual(len(non_empty), 1, "中文多视觉行应合并为一个段落")
        self.assertIn(text, "".join(paras))
        # 纯中文文本，合并后不应出现 ASCII 空格（接续规则对 CJK 侧不补空格）。
        self.assertNotIn(" ", non_empty[0])

    def test_english_cross_line_spacing(self):
        # 英文长段落：跨视觉行处应补一个空格，而非连写或丢失空格。
        en_path = os.path.join(self._tmp.name, "en.pdf")
        _make_en_pdf(en_path)
        self.adapter.convert(
            ConvertPdfToWordRequest(
                pdf_path=en_path, template_path="", output_path=self.out_path
            ),
            lambda c, t: None,
        )
        paras, _ = _collect(self.out_path)
        non_empty = [p for p in paras if p.strip()]
        self.assertEqual(len(non_empty), 1, "英文多视觉行应合并为一个段落")
        joined = non_empty[0]
        self.assertIn("wrapped across", joined)
        self.assertNotIn("wrappedacross", joined)

    def test_invalid_pdf_raises(self):
        bad = os.path.join(self._tmp.name, "bad.pdf")
        with open(bad, "wb") as f:
            f.write(b"not a pdf")
        with self.assertRaises(PdfReadError):
            self.adapter.convert(
                ConvertPdfToWordRequest(
                    pdf_path=bad, template_path="",
                    output_path=self.out_path,
                ),
                lambda c, t: None,
            )


if __name__ == "__main__":
    unittest.main()
