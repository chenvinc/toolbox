"""JSON→Word 试卷服务单元测试（core 纯逻辑，无 Qt、无 GUI）。

用替身（FakeExamDocxWriter / CollectingEmitter）注入端口，验证：
  - 排版常量解析（字号映射、行距预设）
  - JSON 解析（正常 / 空题目 / 非法 JSON）
  - JsonToWordServiceImpl 编排：进度事件 + 完成事件 + 结果
  - DocxExamWriterAdapter 端到端生成可读的题本/解析文档并套用排版
  - 图片：占位符→内联插入（按原始比例）；下载失败→灰色占位框
  - 字体中英拆分（eastAsia / ascii）、解析三段式结构
"""
import base64
import json
import os
import tempfile
import unittest
from pathlib import Path
from typing import Callable, List
from unittest.mock import patch

from shared.contracts import (
    EventType,
    ExamLineSpacingType,
    GenerateExamRequest,
    GenerateExamResult,
)
from core.models.exam_question import ExamImage, ExamQuestion
from core.ports.io import ExamDocxWriter
from core.services._exam_layout import (
    WORD_FONT_SIZE_NAME_TO_PT,
    resolve_exam_line_spacing,
    resolve_font_size_pt,
)
from core.services._exam_parser import parse_exam_json
from core.services.json_to_word_service import JsonToWordServiceImpl
from core.adapters.docx_exam_writer import DocxExamWriterAdapter
from shared.errors import DocumentReadError, NoQuestionsExtracted

from docx import Document
from docx.oxml.ns import qn

# 8x6 红色 PNG（有效），用于本地 file:// 图片插入测试（避免依赖外部网络）。
_PNG_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAgAAAAGCAIAAABxZ0isAAAAEklEQVR4nGP8z4AdMOEQ"
    "p4sEAJiFAQtsMIKbAAAAAElFTkSuQmCC"
)


def _drawing_count(doc: Document) -> int:
    return len(doc.element.body.findall(".//" + qn("w:drawing")))


class CollectingEmitter:
    """实现 EventEmitter 端口，收集所有事件供断言。"""

    def __init__(self):
        self.events = []
        self._handlers = []

    def emit(self, event):
        self.events.append(event)
        for h in self._handlers:
            h(event)

    def on_event(self, handler):
        self._handlers.append(handler)


class FakeExamDocxWriter(ExamDocxWriter):
    """ExamDocxWriter 替身：记录调用与进度回调，返回固定结果。"""

    def __init__(self):
        self.calls = []
        self.progress_calls = []

    def build(self, request, questions, on_progress, image_cache=None):
        self.calls.append((request, list(questions), image_cache))
        if on_progress:
            on_progress(1, 2)
            on_progress(2, 2)
        return GenerateExamResult(
            question_book_path="/tmp/fake_题本.docx",
            analysis_path="/tmp/fake_解析.docx",
            question_count=len(questions),
        )


SAMPLE_JSON = {
    "pageTitle": "测试试卷",
    "questions": [
        {
            "questionNumber": "1.",
            "questionType": "单选题",
            "questionStem": "下列正确的是？",
            "options": {"A": {"text": "甲"}, "B": {"text": "乙"}},
            "correctAnswer": "A",
            "correctRate": "50 %",
            "solution": {"analysis": "因为甲正确。"},
        }
    ],
}


class LayoutResolveTests(unittest.TestCase):
    def test_font_size_map_known(self):
        self.assertEqual(resolve_font_size_pt("五号"), 10.5)
        self.assertEqual(resolve_font_size_pt("小四"), 12.0)

    def test_font_size_unknown_falls_back(self):
        self.assertEqual(resolve_font_size_pt("不存在"), WORD_FONT_SIZE_NAME_TO_PT["五号"])

    def test_line_spacing_presets(self):
        self.assertEqual(
            resolve_exam_line_spacing(ExamLineSpacingType.SINGLE, 1.5), 1.0
        )
        self.assertEqual(
            resolve_exam_line_spacing(ExamLineSpacingType.ONE_HALF, 1.5), 1.5
        )
        self.assertEqual(
            resolve_exam_line_spacing(ExamLineSpacingType.DOUBLE, 1.5), 2.0
        )
        self.assertEqual(
            resolve_exam_line_spacing(ExamLineSpacingType.CUSTOM, 1.8), 1.8
        )


class ParserTests(unittest.TestCase):
    def setUp(self):
        self._dir = tempfile.mkdtemp()

    def _write(self, data) -> str:
        path = os.path.join(self._dir, "exam.json")
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(data, fh, ensure_ascii=False)
        return path

    # ---- 旧 schema 兼容（questionStem / options.text / correctAnswer）----
    def test_parse_ok_legacy(self):
        path = self._write(SAMPLE_JSON)
        questions, title = parse_exam_json(path)
        self.assertEqual(title, "测试试卷")
        self.assertEqual(len(questions), 1)
        q = questions[0]
        self.assertEqual(q.number, "1.")
        self.assertEqual(q.options, {"A": "A. 甲", "B": "B. 乙"})
        self.assertEqual(q.correct_answer, "A")
        self.assertEqual(q.analysis, "因为甲正确。")
        self.assertEqual(q.correct_rate, "50%")

    def test_parse_empty_raises(self):
        path = self._write({"questions": []})
        with self.assertRaises(NoQuestionsExtracted):
            parse_exam_json(path)

    def test_parse_bad_json_raises(self):
        path = os.path.join(self._dir, "bad.json")
        with open(path, "w", encoding="utf-8") as fh:
            fh.write("{ not valid json")
        with self.assertRaises(DocumentReadError):
            parse_exam_json(path)

    # ---- 关键一：HTML 剥离（标签 / Angular 属性 / 内联样式 / <img>→[IMGn]）----
    def test_html_strip_removes_tags_keeps_img_placeholder(self):
        html = (
            '<div _ngcontent-ng-c1=""><p>求 <b>面积</b>？</p>'
            '<p><img width="100" src="//cdn.example.com/p1.png?x=1"></p></div>'
        )
        data = {
            "pageTitle": "T",
            "questions": [{
                "questionNumber": "1.", "questionType": "单选",
                "contentHtml": html,
                "options": {"A": {"text": "对"}},
                "correctAnswer": "A", "correctRate": "60 %",
                "solution": {"analysis": "略"},
                "images": [{"index": 1, "src": "https://cdn.example.com/p1.png?x=1",
                            "role": "stem", "isTex": False}],
            }],
        }
        q = parse_exam_json(self._write(data))[0][0]
        self.assertNotIn("<", q.stem)
        self.assertNotIn("_ngcontent", q.stem)
        self.assertNotIn("style", q.stem)
        self.assertIn("[IMG1]", q.stem)
        self.assertIn("面积", q.stem)

    # ---- 关键二：图片索引匹配 + role / isTex ----
    def test_image_index_role_and_istex(self):
        data = {
            "pageTitle": "T",
            "questions": [{
                "questionNumber": "1.", "questionType": "单选",
                "contentHtml": '<p>图：<img src="//cdn/x/stem.png"></p>',
                "options": {"A": {"text": "a"}},
                "correctAnswer": "A", "correctRate": "10 %",
                "solution": {"analysis": '<p>解析图：<img src="//cdn/x/sol.png"></p>'},
                "images": [
                    {"index": 1, "src": "https://cdn/x/stem.png", "role": "stem", "isTex": False},
                    {"index": 2, "src": "https://cdn/x/sol.png", "role": "solution", "isTex": True},
                ],
            }],
        }
        q = parse_exam_json(self._write(data))[0][0]
        self.assertEqual(len(q.images), 2)
        self.assertEqual(q.images[0].index, 1)
        self.assertEqual(q.images[0].role, "stem")
        self.assertFalse(q.images[0].is_tex)
        self.assertEqual(q.images[1].index, 2)
        self.assertEqual(q.images[1].role, "solution")
        self.assertTrue(q.images[1].is_tex)
        # 占位符 [IMGn] 与 images.index 对应
        self.assertIn("[IMG1]", q.stem)
        self.assertIn("[IMG2]", q.analysis)

    # ---- 关键三：选项格式化（忽略 HTML，跳过空选项）----
    def test_option_formatting_strips_html_and_skips_empty(self):
        data = {
            "pageTitle": "T",
            "questions": [{
                "questionNumber": "1.", "questionType": "单选",
                "contentHtml": "<p>题干</p>",
                "options": {
                    "A": {"text": "186", "html": '<input type="radio">186'},
                    "B": {"text": "187", "html": '<input type="radio">187'},
                    "C": {"text": "", "html": "<input>"},
                    "D": {"text": "189", "html": '<input type="radio">189'},
                },
                "correctAnswer": "D", "correctRate": "38 %",
                "solution": {"analysis": "x"},
            }],
        }
        q = parse_exam_json(self._write(data))[0][0]
        self.assertEqual(q.options, {"A": "A. 186", "B": "B. 187", "D": "D. 189"})
        self.assertNotIn("C", q.options)  # 空选项跳过
        self.assertNotIn("<", q.options["A"])

    # ---- 正确率归一化 ----
    def test_correct_rate_normalization(self):
        base = {
            "pageTitle": "T",
            "questions": [{
                "questionNumber": "1.", "questionType": "单选",
                "contentHtml": "<p>x</p>",
                "options": {"A": {"text": "a"}},
                "correctAnswer": "A", "correctRate": "38 %",
                "solution": {"analysis": "y"},
            }],
        }
        q = parse_exam_json(self._write(base))[0][0]
        self.assertEqual(q.correct_rate, "38%")

        base["questions"][0]["correctRate"] = 0.68
        q2 = parse_exam_json(self._write(base))[0][0]
        self.assertEqual(q2.correct_rate, "68%")

    # ---- 真实样本端到端验证 ----
    def test_parse_real_sample(self):
        sample = (Path(__file__).resolve().parents[3] / "docs" /
                  "专项智能练习（数量关系）_2026-07-17T06-01-08-897.json")
        if not sample.exists():
            self.skipTest(f"样本文件缺失：{sample}")
        questions, title = parse_exam_json(str(sample))
        self.assertEqual(title, "专项智能练习（数量关系）")
        self.assertEqual(len(questions), 3)

        q0 = questions[0]
        # 题干：纯文本 + [IMG1]，无 HTML / Angular 属性
        self.assertIn("[IMG1]", q0.stem)
        self.assertNotIn("<", q0.stem)
        self.assertNotIn("_ngcontent", q0.stem)
        # 选项格式化（忽略 radio 等 HTML）
        self.assertEqual(q0.options, {
            "A": "A. 186", "B": "B. 187", "C": "C. 188", "D": "D. 189",
        })
        # 答案 / 正确率
        self.assertEqual(q0.correct_answer, "D")
        self.assertEqual(q0.correct_rate, "38%")
        # 解析：含 [IMG2]..[IMG5]，无 HTML
        self.assertIn("[IMG2]", q0.analysis)
        self.assertIn("[IMG5]", q0.analysis)
        self.assertNotIn("<", q0.analysis)
        # 图片映射（index / role / isTex）
        self.assertEqual(len(q0.images), 5)
        self.assertEqual(q0.images[0].index, 1)
        self.assertEqual(q0.images[0].role, "stem")
        self.assertFalse(q0.images[0].is_tex)
        self.assertEqual(q0.images[1].index, 2)
        self.assertEqual(q0.images[1].role, "solution")
        self.assertTrue(q0.images[1].is_tex)


class ServiceTests(unittest.TestCase):
    def setUp(self):
        self._dir = tempfile.mkdtemp()

    def _write_json(self) -> str:
        path = os.path.join(self._dir, "exam.json")
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(SAMPLE_JSON, fh, ensure_ascii=False)
        return path

    def test_generate_emits_progress_and_completed(self):
        emitter = CollectingEmitter()
        writer = FakeExamDocxWriter()
        svc = JsonToWordServiceImpl(writer, emitter)

        # 服务会先解析 JSON（再调用 writer），故需提供可解析的临时文件
        req = GenerateExamRequest(input_path=self._write_json())
        result = svc.generate(req)

        self.assertIsInstance(result, GenerateExamResult)
        self.assertEqual(result.question_count, 1)  # SAMPLE_JSON 含 1 题，writer 回显题数
        types = [e.type for e in emitter.events]
        self.assertIn(EventType.EXAM_PROGRESS, types)
        self.assertIn(EventType.EXAM_COMPLETED, types)
        # 完成事件为最后一个
        self.assertEqual(emitter.events[-1].type, EventType.EXAM_COMPLETED)


class AdapterE2ETests(unittest.TestCase):
    """验证适配器真正写出可读的 Word 文档并套用排版设置。"""

    def setUp(self):
        self._dir = tempfile.mkdtemp()

    def _write_json(self) -> str:
        path = os.path.join(self._dir, "exam.json")
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(SAMPLE_JSON, fh, ensure_ascii=False)
        return path

    def _read_paras(self, path: str) -> List[str]:
        doc = Document(path)
        return [p.text for p in doc.paragraphs if p.text.strip()]

    def test_adapter_writes_book_and_analysis(self):
        json_path = self._write_json()
        req = GenerateExamRequest(input_path=json_path, output_dir=self._dir)
        writer = DocxExamWriterAdapter()
        result = writer.build(req, parse_exam_json(json_path)[0], on_progress=lambda *_: None)

        self.assertTrue(os.path.exists(result.question_book_path))
        self.assertTrue(os.path.exists(result.analysis_path))

        book = self._read_paras(result.question_book_path)
        self.assertTrue(any("下列正确的是？" in p for p in book))
        self.assertTrue(any(p.startswith("A. ") for p in book))
        self.assertFalse(any("【答案】" in p for p in book))  # 题本不含答案

        analysis = self._read_paras(result.analysis_path)
        self.assertTrue(any("【答案】A" in p for p in analysis))
        self.assertTrue(any("【解析】" in p for p in analysis))

    def test_adapter_applies_first_line_indent(self):
        json_path = self._write_json()
        req = GenerateExamRequest(
            input_path=json_path, output_dir=self._dir,
            line_spacing_type=ExamLineSpacingType.DOUBLE,
            first_line_indent=True,
        )
        writer = DocxExamWriterAdapter()
        result = writer.build(req, parse_exam_json(json_path)[0], on_progress=lambda *_: None)
        doc = Document(result.question_book_path)
        # 题干段落应带首行缩进且行距为 2.0 倍
        stem_para = next(p for p in doc.paragraphs if "下列正确的是？" in p.text)
        self.assertIsNotNone(stem_para.paragraph_format.first_line_indent)
        self.assertAlmostEqual(float(stem_para.paragraph_format.line_spacing), 2.0, places=3)


class AdapterImageTests(unittest.TestCase):
    """验证图片占位符→内联插入（失败→灰色占位框）与字体中英拆分。"""

    def setUp(self):
        self._dir = tempfile.mkdtemp()
        self._png = os.path.join(self._dir, "px.png")
        with open(self._png, "wb") as fh:
            fh.write(base64.b64decode(_PNG_B64))
        self._url = "file://" + self._png

    def _question(self) -> List[ExamQuestion]:
        return [ExamQuestion(
            number="1.",
            question_type="单选",
            stem="看图：\n[IMG1]",
            options={"A": "A. 对", "B": "B. 错"},
            correct_answer="A",
            correct_rate="50%",
            analysis="[IMG1] 如图所示。",
            images=[ExamImage(index=1, src=self._url, role="stem", is_tex=False)],
        )]

    def _build(self, questions: List[ExamQuestion], **kw) -> GenerateExamResult:
        json_path = os.path.join(self._dir, "exam.json")
        with open(json_path, "w", encoding="utf-8") as fh:
            json.dump({"pageTitle": "图卷", "questions": []}, fh, ensure_ascii=False)
        req = GenerateExamRequest(input_path=json_path, output_dir=self._dir, **kw)
        writer = DocxExamWriterAdapter()
        return writer.build(req, questions, on_progress=lambda *_: None)

    def test_image_inserted_for_placeholder(self):
        res = self._build(self._question())
        doc = Document(res.question_book_path)
        self.assertEqual(_drawing_count(doc), 1)  # 题本 1 张图
        # 占位符文本不应残留为字面量
        self.assertFalse(any("[IMG1]" in p.text for p in doc.paragraphs))
        # 解析也应含该图（同 index）
        self.assertEqual(_drawing_count(Document(res.analysis_path)), 1)

    def test_image_failure_inserts_placeholder(self):
        # 强制下载失败：不抛异常，插入灰色占位框文字
        with patch(
            "core.adapters.docx_exam_writer._fetch_image_bytes", return_value=None
        ):
            res = self._build(self._question())
        doc = Document(res.question_book_path)
        self.assertEqual(_drawing_count(doc), 0)
        self.assertTrue(any("[图片加载失败:" in p.text for p in doc.paragraphs))

    def test_font_split_cjk_latin(self):
        res = self._build(self._question(), font_name="宋体/Times New Roman")
        doc = Document(res.question_book_path)
        run = doc.paragraphs[0].runs[0]
        rpr = run._element.find(qn("w:rPr"))
        rf = rpr.find(qn("w:rFonts")) if rpr is not None else None
        self.assertIsNotNone(rf)
        self.assertEqual(rf.get(qn("w:eastAsia")), "宋体")
        self.assertEqual(rf.get(qn("w:ascii")), "Times New Roman")

    def test_analysis_three_segment_structure(self):
        res = self._build(self._question())
        texts = [p.text for p in Document(res.analysis_path).paragraphs if p.text.strip()]
        # 第一段：带题号的【答案】
        self.assertTrue(texts[0].startswith("1.【答案】A"))
        # 第二段：【正确率】（不带题号）
        self.assertTrue(texts[1].startswith("【正确率】"))
        self.assertFalse(texts[1].startswith("1."))
        # 第三段起：【解析】
        self.assertTrue(any(t.startswith("【解析】") for t in texts[2:]))


if __name__ == "__main__":
    unittest.main()
