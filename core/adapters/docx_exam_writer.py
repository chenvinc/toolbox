"""试卷（题本 + 解析）Word 文档写适配器（封装 python-docx）。

实现 core/ports/io.ExamDocxWriter。把原「Word 试卷生成」的纯生成逻辑迁至此，
由 JsonToWordServiceImpl 通过端口调用；外部（测试）可用 FakeExamDocxWriter 替身。
排版设置（字体 / 字号 / 行间距 / 首行缩进）全部来自 GenerateExamRequest，
对题本文档与解析文档同时生效。

图片处理：对题干 / 解析中的 ``[IMGn]`` 占位符，按题内 images 映射下载并内联插入，
保持原图宽高比；下载失败则就地插入灰色占位框（含 URL），不中断整体流程。
"""
from __future__ import annotations

import io
import os
import re
from typing import Callable, Dict, List, Optional, Tuple, TYPE_CHECKING

from docx import Document
from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section
from docx.shared import Emu, Pt
from docx.image.image import Image as _DocxImage  # 仅用于读取原图尺寸（按原始比例插入）

from shared.contracts import (
    GenerateExamRequest,
    GenerateExamResult,
)
from shared.errors import OutputWriteError
from core.models.exam_question import ExamImage, ExamQuestion
from core.ports.io import ExamDocxWriter
from core.services._exam_image_fetch import _fetch_image_bytes
from core.services._exam_layout import resolve_exam_line_spacing, resolve_font_size_pt
from core.services._exam_parser import read_page_title

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

# 中文字号下「2 字符」首行缩进的基准（EMU）。运行时按字号 ×2 计算。
# 中文 Word 默认页边距（与参考 docx 一致）：上/下 1 英寸、左/右 1.25 英寸。
_MARGIN_TOP_BOTTOM_EMU = 914400    # 1 inch
_MARGIN_LEFT_RIGHT_EMU = 1143000   # 1.25 inch
# 题与题之间留的段前间距（适当段间距）。
_QUESTION_GAP_PT = 6
# 图片加载失败时占位框的底色（浅灰）。
_PLACEHOLDER_FILL = "D9D9D9"

# 题干 / 解析文本中的图片占位符： [IMGn]
_IMG_PLACEHOLDER_RE = re.compile(r"\[IMG(\d+)\]")


def _sanitize_filename(name: str) -> str:
    """去除文件名非法字符并限长，用于把 pageTitle 安全地用作输出文件名。"""
    cleaned = re.sub(r'[\\/:*?"<>|\r\n\t]', "", (name or "").strip())
    return cleaned[:80]


def _split_font(font_name: str) -> Tuple[str, str]:
    """把 ``CJK/Latin`` 组合字体名拆成 (中文, 英文)；无斜杠时中英同体。"""
    if "/" in font_name:
        cjk, latin = font_name.split("/", 1)
        return cjk.strip(), latin.strip()
    return font_name.strip(), font_name.strip()


def _set_run_font(run: "Run", cjk_font: str, latin_font: str, font_size_pt: float) -> None:
    """为文本 run 设置中英分字体与字号。

    中文（eastAsia）使用用户指定中文字体；英文 / 数字（ascii / hAnsi）使用
    Times New Roman 或用户指定的拉丁字体。
    """
    run.font.name = latin_font
    run.font.size = Pt(font_size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cjk_font)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)


def _apply_paragraph_format(
    paragraph: "Paragraph",
    *,
    cjk_font: str,
    latin_font: str,
    font_size_pt: float,
    line_spacing: float,
    first_line_indent: bool,
) -> None:
    """统一设置段落的字体、行距与首行缩进（2 字符）。"""
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    # 1 字符 ≈ 1 倍字号；首行缩进 2 字符 = 2 × 字号（pt）。
    # 纯图片段落（无文字）不缩进，避免图片被右移 2 字符。
    has_drawing = bool(paragraph._p.findall(".//" + qn("w:drawing")))
    has_text = bool(paragraph.text.strip())
    if first_line_indent and not (has_drawing and not has_text):
        fmt.first_line_indent = Pt(font_size_pt * 2)
    else:
        fmt.first_line_indent = Pt(0)
    for run in paragraph.runs:
        _set_run_font(run, cjk_font, latin_font, font_size_pt)


def _shade_paragraph(paragraph: "Paragraph", fill: str) -> None:
    """给段落加底纹（灰色占位框用）。"""
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def _set_margins(section: Section) -> None:
    """把节边距设为中文 Word 默认（与参考 docx 对齐）。"""
    section.left_margin = Emu(_MARGIN_LEFT_RIGHT_EMU)
    section.right_margin = Emu(_MARGIN_LEFT_RIGHT_EMU)
    section.top_margin = Emu(_MARGIN_TOP_BOTTOM_EMU)
    section.bottom_margin = Emu(_MARGIN_TOP_BOTTOM_EMU)


def _usable_width_emu(section: Section) -> int:
    """文本区域可用宽度（页宽 - 左右边距）。"""
    page_w = int(section.page_width) if section.page_width is not None else 7772400
    return page_w - _MARGIN_LEFT_RIGHT_EMU * 2


def _fit_image_size(native_w: int, native_h: int, usable: int) -> Tuple[Emu, Emu]:
    """按原始宽高比计算显示尺寸：不超过可用宽度（避免小图被放大、大图溢出）。"""
    if native_w <= usable:
        return Emu(native_w), Emu(native_h)
    ratio = native_h / native_w
    return Emu(usable), Emu(int(usable * ratio))


def _insert_image_or_placeholder(
    paragraph: "Paragraph",
    img: ExamImage,
    cache: Dict[str, Optional[bytes]],
    failed: List[str],
    cjk_font: str,
    latin_font: str,
    font_size_pt: float,
    usable_emu: int,
) -> None:
    """在段落内联插入图片；失败则插入灰色占位框（含 URL）并记录失败。

    失败分两类，都必须记入 ``failed`` 以便上层如实汇报：
      * 下载失败（``cache[src]`` 为 None）；
      * 下载成功但字节不可用 ``_DocxImage.from_blob`` 无法识别（如源站返回
        加密/非图片内容、SVG 等 python-docx 不支持的格式）。
    """
    data = _fetch_image_bytes(img.src, cache)
    if data is None:
        run = paragraph.add_run(f"[图片加载失败: {img.src}]")
        _set_run_font(run, cjk_font, latin_font, font_size_pt)
        _shade_paragraph(paragraph, _PLACEHOLDER_FILL)
        if img.src not in failed:
            failed.append(img.src)
        return
    try:
        native = _DocxImage.from_blob(data)
        nw = int(native.width)
        nh = int(native.height)
        w, h = _fit_image_size(nw, nh, usable_emu)
        run = paragraph.add_run("")
        run.add_picture(io.BytesIO(data), width=w, height=h)
    except Exception:
        run = paragraph.add_run(f"[图片加载失败: {img.src}]")
        _set_run_font(run, cjk_font, latin_font, font_size_pt)
        _shade_paragraph(paragraph, _PLACEHOLDER_FILL)
        if img.src not in failed:
            failed.append(img.src)


def _add_text_with_images(
    paragraph: "Paragraph",
    text: str,
    img_by_index: Dict[int, ExamImage],
    cache: Dict[str, Optional[bytes]],
    failed: List[str],
    cjk_font: str,
    latin_font: str,
    font_size_pt: float,
    usable_emu: int,
) -> None:
    """把文本写入段落，并把 ``[IMGn]`` 替换为内联图片（保留文字与图片混排）。"""
    last = 0
    for m in _IMG_PLACEHOLDER_RE.finditer(text):
        before = text[last:m.start()]
        if before:
            run = paragraph.add_run(before)
            _set_run_font(run, cjk_font, latin_font, font_size_pt)
        idx = int(m.group(1))
        img = img_by_index.get(idx)
        if img is not None:
            _insert_image_or_placeholder(
                paragraph, img, cache, failed, cjk_font, latin_font, font_size_pt, usable_emu
            )
        # 未知占位符：直接丢弃，不写入任何内容。
        last = m.end()
    after = text[last:]
    if after:
        run = paragraph.add_run(after)
        _set_run_font(run, cjk_font, latin_font, font_size_pt)


class DocxExamWriterAdapter(ExamDocxWriter):
    """基于 python-docx 生成题本与解析两份 Word 文档。"""

    def build(
        self,
        request: GenerateExamRequest,
        questions: List[ExamQuestion],
        on_progress: Callable[[int, int], None],
        image_cache: Optional[Dict[str, Optional[bytes]]] = None,
    ) -> GenerateExamResult:
        out_dir = request.output_dir.strip() or os.path.dirname(request.input_path)
        try:
            os.makedirs(out_dir, exist_ok=True)
        except OSError as exc:
            raise OutputWriteError(
                f"无法创建输出目录：{out_dir}（{exc}）", output_dir=out_dir
            ) from exc

        # 输出文件名取自 pageTitle（用户要求），非法字符去除后作为安全文件名；
        # 回退到输入 JSON 文件名，避免空标题导致文件名异常。
        page_title = read_page_title(request.input_path)
        safe_title = _sanitize_filename(page_title) or os.path.splitext(
            os.path.basename(request.input_path)
        )[0]
        question_book_path = os.path.join(out_dir, f"{safe_title}_题本.docx")
        analysis_path = os.path.join(out_dir, f"{safe_title}_解析.docx")

        cjk_font, latin_font = _split_font(request.font_name)
        font_size_pt = resolve_font_size_pt(request.font_size_name)
        line_spacing = resolve_exam_line_spacing(
            request.line_spacing_type, request.line_spacing_value
        )
        indent = request.first_line_indent

        # 预下载缓存由服务层注入（已并发下载）；未注入时建空缓存走惰性回退下载。
        cache: Dict[str, Optional[bytes]] = image_cache if image_cache is not None else {}

        # 收集「插入阶段」失败（含下载失败与下载成功但字节不可识别两类），
        # 用于如实汇报，避免 UI 误报「全部成功」。
        failed_book: List[str] = []
        failed_analysis: List[str] = []
        try:
            self._build_question_book(
                question_book_path, questions,
                cjk_font, latin_font, font_size_pt, line_spacing, indent, cache, failed_book,
            )
            on_progress(1, 2)

            self._build_analysis(
                analysis_path, questions,
                cjk_font, latin_font, font_size_pt, line_spacing, indent, cache, failed_analysis,
            )
            on_progress(2, 2)
        except OSError as exc:
            raise OutputWriteError(
                f"写入文件失败（输出目录可能无写入权限）：{exc}", output_dir=out_dir
            ) from exc

        # 失败图片 = 插入阶段记录的失败 ∪ 下载阶段记为 None 的 URL（去重保序）。
        inserted_failures = list(dict.fromkeys(failed_book + failed_analysis))
        download_none = [src for src, data in cache.items() if data is None]
        failed_images = list(dict.fromkeys(inserted_failures + download_none))

        return GenerateExamResult(
            question_book_path=question_book_path,
            analysis_path=analysis_path,
            question_count=len(questions),
            failed_images=failed_images,
        )

    # ── 题本：题号 + 题干（含图）+ 选项（不含答案） ──
    def _build_question_book(
        self,
        out_path: str,
        questions: List[ExamQuestion],
        cjk_font: str,
        latin_font: str,
        font_size_pt: float,
        line_spacing: float,
        first_line_indent: bool,
        cache: Dict[str, Optional[bytes]],
        failed: List[str],
    ) -> None:
        doc: _Document = Document()
        _set_margins(doc.sections[0])
        usable = _usable_width_emu(doc.sections[0])

        for qi, q in enumerate(questions):
            img_by_index = {im.index: im for im in q.images}
            stem_lines = q.stem.split("\n")

            # 第一行题干带题号；行内的 [IMGn] 内联为图片。
            first_p = doc.add_paragraph()
            _add_text_with_images(
                first_p, f"{q.number}{stem_lines[0]}", img_by_index, cache, failed,
                cjk_font, latin_font, font_size_pt, usable,
            )
            _apply_paragraph_format(
                first_p, cjk_font=cjk_font, latin_font=latin_font,
                font_size_pt=font_size_pt, line_spacing=line_spacing,
                first_line_indent=first_line_indent,
            )
            if qi > 0:
                first_p.paragraph_format.space_before = Pt(_QUESTION_GAP_PT)

            # 题干中换行后的剩余行（如图片独占一行）各自成段。
            for extra in stem_lines[1:]:
                extra_p = doc.add_paragraph()
                _add_text_with_images(
                    extra_p, extra, img_by_index, cache, failed,
                    cjk_font, latin_font, font_size_pt, usable,
                )
                _apply_paragraph_format(
                    extra_p, cjk_font=cjk_font, latin_font=latin_font,
                    font_size_pt=font_size_pt, line_spacing=line_spacing,
                    first_line_indent=first_line_indent,
                )

            # 选项：A. xxx 各占一行。
            for key in sorted(q.options):
                opt_p = doc.add_paragraph()
                _add_text_with_images(
                    opt_p, q.options[key], img_by_index, cache, failed,
                    cjk_font, latin_font, font_size_pt, usable,
                )
                _apply_paragraph_format(
                    opt_p, cjk_font=cjk_font, latin_font=latin_font,
                    font_size_pt=font_size_pt, line_spacing=line_spacing,
                    first_line_indent=first_line_indent,
                )
        doc.save(out_path)

    # ── 解析：题号 + 【答案】 / 【正确率】 / 【解析】+ 多段解析（含图） ──
    def _build_analysis(
        self,
        out_path: str,
        questions: List[ExamQuestion],
        cjk_font: str,
        latin_font: str,
        font_size_pt: float,
        line_spacing: float,
        first_line_indent: bool,
        cache: Dict[str, Optional[bytes]],
        failed: List[str],
    ) -> None:
        doc: _Document = Document()
        _set_margins(doc.sections[0])
        usable = _usable_width_emu(doc.sections[0])

        for qi, q in enumerate(questions):
            img_by_index = {im.index: im for im in q.images}

            # 第一段：【答案】X（带题号，与题本题号格式一致）。
            answer_p = doc.add_paragraph()
            _add_text_with_images(
                answer_p, f"{q.number}【答案】{q.correct_answer}", img_by_index, cache, failed,
                cjk_font, latin_font, font_size_pt, usable,
            )
            _apply_paragraph_format(
                answer_p, cjk_font=cjk_font, latin_font=latin_font,
                font_size_pt=font_size_pt, line_spacing=line_spacing,
                first_line_indent=first_line_indent,
            )
            if qi > 0:
                answer_p.paragraph_format.space_before = Pt(_QUESTION_GAP_PT)

            # 第二段：【正确率】XX%（不带题号，与参考 docx 一致）。
            rate_p = doc.add_paragraph()
            _add_text_with_images(
                rate_p, f"【正确率】{q.correct_rate}", img_by_index, cache, failed,
                cjk_font, latin_font, font_size_pt, usable,
            )
            _apply_paragraph_format(
                rate_p, cjk_font=cjk_font, latin_font=latin_font,
                font_size_pt=font_size_pt, line_spacing=line_spacing,
                first_line_indent=first_line_indent,
            )

            # 第三段起：【解析】+ 解析正文；按换行拆多段，[IMGn] 内联为图片。
            analysis_lines = q.analysis.split("\n")
            for ai, seg in enumerate(analysis_lines):
                ana_p = doc.add_paragraph()
                prefix = "【解析】" if ai == 0 else ""
                _add_text_with_images(
                    ana_p, prefix + seg, img_by_index, cache, failed,
                    cjk_font, latin_font, font_size_pt, usable,
                )
                _apply_paragraph_format(
                    ana_p, cjk_font=cjk_font, latin_font=latin_font,
                    font_size_pt=font_size_pt, line_spacing=line_spacing,
                    first_line_indent=first_line_indent,
                )
        doc.save(out_path)
