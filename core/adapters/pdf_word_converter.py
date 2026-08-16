"""PDF → Word 转换适配器（封装 PyMuPDF + python-docx）。

实现 core/ports/io.PdfWordConverter。把 PDF 文字按阅读顺序重排为「语义段落」，
而非「逐视觉行」。段落重建两阶段：

1. 块内合并（block-internal）：PyMuPDF 的 ``block`` 本身是连贯文本区，块内多行
   只是同一段落的视觉换行，因此块内所有行无条件合并为一个候选段落。
2. 块间智能合并（smart merge）：相邻候选段落若「垂直间距 ≤ 1.3× 行高中位数」且
   「无首行缩进跳变（≥0.5× 字号）」且「字号未显著突变」，则视为同一段落继续合并。
   这可还原被 PyMuPDF 拆成多块的跨块段落，同时避免把标题/不同段落误并。

行间接续规则（connection rule）：上一行末字符或下一行首字符任一侧为 CJK/中文标点
则不加空格；两侧皆为拉丁字母/数字才补一个空格——以此修复英文跨行连写与中文
「无词间空格」导致的错位。首版纯文字，不提取图片（与 PdfSlideConverter
的「不栅格化」哲学一致）。

字体归一 / 粗斜体判定复用 PdfSlideConverter 的纯函数（clean_font / is_bold /
is_italic），避免重复实现。
"""
from __future__ import annotations

import statistics
from dataclasses import dataclass, field
from typing import Callable, List, Optional

import fitz  # PyMuPDF
from docx import Document
from docx.document import Document as _Document
from docx.shared import Pt
from docx.text.run import Run

from core.adapters.pdf_slide_converter import (
    clean_font,
    is_bold,
    is_italic,
)
from core.ports.io import PdfWordConverter
from shared.contracts import ConvertPdfToWordRequest, ConvertPdfToWordResult
from shared.errors import PdfReadError, TemplateInvalidError


# ---------------------------------------------------------------- 数据结构


@dataclass
class _Span:
    """一个文本 span 的最小化表示。"""

    text: str
    font: Optional[str]
    size: float
    color: Optional[int]
    flags: int


@dataclass
class _Line:
    """一行（视觉换行）的几何与内容。"""

    y0: float
    y1: float
    x0: float
    spans: List[_Span] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "".join(s.text for s in self.spans)

    @property
    def start_char(self) -> Optional[str]:
        t = self.text
        return t[0] if t else None

    @property
    def end_char(self) -> Optional[str]:
        t = self.text
        return t[-1] if t else None

    @property
    def size(self) -> float:
        for s in self.spans:
            if s.size:
                return s.size
        return 12.0


# ---------------------------------------------------------------- 模板清空


def _clear_body(doc: _Document) -> None:
    """清空文档正文（段落与表格），保留节 / 样式 / 页眉页脚定义。

    套用模板时据此抹掉模板自带示例内容，仅留样式骨架。
    """
    body = doc.element.body
    for p in list(doc.paragraphs):
        el = p._p
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    for tbl in list(doc.tables):
        el = tbl._tbl
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


# ---------------------------------------------------------------- 接续规则


def _is_cjk(ch: Optional[str]) -> bool:
    """判断字符是否为 CJK 或中文标点（接续规则用于决定补不补空格）。"""
    if not ch:
        return False
    o = ord(ch)
    return (
        0x3000 <= o <= 0x303F  # CJK 符号与标点
        or 0x3400 <= o <= 0x4DBF  # CJK 扩展 A
        or 0x4E00 <= o <= 0x9FFF  # CJK 基本汉字
        or 0xF900 <= o <= 0xFAFF  # CJK 兼容汉字
        or 0xFF00 <= o <= 0xFFEF  # 全角字符（含全角标点）
        or 0x3040 <= o <= 0x30FF  # 假名
    )


def _join_sep(end_char: Optional[str], start_char: Optional[str]) -> str:
    """返回两行接续处应插入的分隔串：CJK 相关侧为空，拉丁-拉丁补一个空格。"""
    if not end_char or not start_char:
        return ""
    if _is_cjk(end_char) or _is_cjk(start_char):
        return ""  # 中文/中文标点任一侧 → 无词间空格
    return " "  # 两侧皆拉丁字母/数字 → 补空格，修复英文跨行连写


# ---------------------------------------------------------------- 合并判定


def _should_merge(prev: List[_Line], nxt: List[_Line], median_h: float) -> bool:
    """相邻候选段落是否应合并为同一语义段落。

    条件（全部满足才合并）：
      - 垂直间距（上一块末行底 → 下一块首行顶）≤ 1.3× 行高中位数；
      - 无首行缩进跳变：下一块首行 x0 − 上一块首行 x0 < 0.5× 字号；
      - 字号未显著突变（比值差 ≤ 25%），避免把标题并入正文。
    """
    prev_last = prev[-1]
    prev_first = prev[0]
    nxt_first = nxt[0]

    gap = nxt_first.y0 - prev_last.y1
    if gap < 0:
        gap = 0
    if gap > 1.3 * median_h:
        return False

    fontsize = max(nxt_first.size, 1.0)
    indent = nxt_first.x0 - prev_first.x0
    if indent >= 0.5 * fontsize:
        return False

    # 字号显著不同（如标题 vs 正文）则视为不同段落
    bigger = max(prev_first.size, nxt_first.size)
    if bigger > 0 and abs(prev_first.size - nxt_first.size) / bigger > 0.25:
        return False

    return True


# ---------------------------------------------------------------- run 级格式


def _set_run_format(
    run: Run,
    raw_font: Optional[str],
    size: float,
    color: Optional[int],
    flags: int,
) -> None:
    """为一个 Word run 设置字体 / 字号 / 颜色 / 粗斜体。

    同时写入 ``w:rFonts`` 的 ``eastAsia``，确保中文按 PDF 所用中文字体渲染。
    """
    from docx.shared import RGBColor  # docx.dml.color 不显式导出 RGBColor（mypy no_implicit_reexport）
    from docx.oxml.ns import qn

    name = clean_font(raw_font)
    run.font.name = name
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    if color is not None:
        c = int(color) & 0xFFFFFF
        run.font.color.rgb = RGBColor(  # type: ignore[no-untyped-call]
            (c >> 16) & 0xFF, (c >> 8) & 0xFF, c & 0xFF
        )
    if is_bold(raw_font, flags):
        run.font.bold = True
    if is_italic(raw_font, flags):
        run.font.italic = True


def _write_paragraph(doc: _Document, lines: List[_Line]) -> int:
    """把一个语义段落（若干视觉行）写入 docx，返回写入的 run 数。"""
    para = doc.add_paragraph()
    runs = 0
    prev: Optional[_Line] = None
    for ln in lines:
        if prev is not None:
            sep = _join_sep(prev.end_char, ln.start_char)
            if sep:
                # 空格 run 借用下一行首个 span 的格式（空格本身不可见，仅占位）
                fs = ln.spans[0]
                run = para.add_run(sep)
                _set_run_format(run, fs.font, fs.size, fs.color, fs.flags)
                runs += 1
        for sp in ln.spans:
            run = para.add_run(sp.text)
            _set_run_format(run, sp.font, sp.size, sp.color, sp.flags)
            runs += 1
        prev = ln
    return runs


# ---------------------------------------------------------------- 适配器


class PdfWordConverterAdapter(PdfWordConverter):
    """按阅读顺序把 PDF 文字逐页重建为 .docx 可编辑段落。"""

    def convert(
        self,
        request: ConvertPdfToWordRequest,
        on_progress: Callable[[int, int], None],
    ) -> ConvertPdfToWordResult:
        # 基底文档：模板复用其样式，否则空白默认样式。
        if request.template_path:
            try:
                doc = Document(request.template_path)
            except Exception as exc:
                raise TemplateInvalidError(f"模板打开失败：{exc}") from exc
            _clear_body(doc)
        else:
            doc = Document()
            _clear_body(doc)

        try:
            src = fitz.open(request.pdf_path)
        except Exception as exc:
            raise PdfReadError(f"PDF 打开失败：{exc}") from exc

        paragraphs = 0
        runs = 0
        empty_pages: List[int] = []

        try:
            total = int(src.page_count)
            for pi in range(total):
                page = src[pi]
                td = page.get_text("dict")

                # 阅读顺序：块按 y0、行内按 x0 排序（近似单栏/顺序版式）。
                blocks = [
                    b for b in td.get("blocks", []) if b.get("type") == 0
                ]
                blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))

                # 提取每块的视觉行，并收集行高用于中位数。
                block_lines: List[List[_Line]] = []
                all_heights: List[float] = []
                for blk in blocks:
                    raw_lines = sorted(
                        blk.get("lines", []), key=lambda ln: ln["bbox"][1]
                    )
                    ls: List[_Line] = []
                    for line in raw_lines:
                        spans = [
                            s for s in line.get("spans", []) if s.get("text")
                        ]
                        if not spans:
                            continue
                        spans.sort(key=lambda s: s["bbox"][0])
                        sl = _Line(
                            y0=line["bbox"][1],
                            y1=line["bbox"][3],
                            x0=line["bbox"][0],
                            spans=[
                                _Span(
                                    text=s["text"],
                                    font=s.get("font"),
                                    size=s.get("size") or 12,
                                    color=s.get("color"),
                                    flags=s.get("flags", 0),
                                )
                                for s in spans
                            ],
                        )
                        ls.append(sl)
                        all_heights.append(sl.y1 - sl.y0)
                    if ls:
                        block_lines.append(ls)

                page_has_text = bool(block_lines)
                if page_has_text:
                    median_h = (
                        statistics.median(all_heights)
                        if all_heights
                        else 12.0
                    )
                    # 每个块先作为独立候选段落。
                    candidates: List[List[_Line]] = list(block_lines)
                    # 跨块智能合并。
                    merged: List[List[_Line]] = []
                    for cand in candidates:
                        if merged and _should_merge(merged[-1], cand, median_h):
                            merged[-1].extend(cand)
                        else:
                            merged.append(cand)
                    # 落盘。
                    for para_lines in merged:
                        runs += _write_paragraph(doc, para_lines)
                        paragraphs += 1

                if not page_has_text:
                    empty_pages.append(pi + 1)  # 纯图片页(如封面)，源 PDF 无文字

                on_progress(pi + 1, total)

            doc.save(request.output_path)
        finally:
            src.close()

        return ConvertPdfToWordResult(
            output_path=request.output_path,
            page_count=total,
            paragraph_count=paragraphs,
            run_count=runs,
            empty_pages=empty_pages,
        )
