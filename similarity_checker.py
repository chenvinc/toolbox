"""相似题目检测工具 — 检测主文档与多个副文档之间的题目重复情况。"""

import os
import re
import subprocess
import sys
from collections import Counter
from difflib import SequenceMatcher

from docx import Document
from PySide6.QtWidgets import (
    QApplication, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QFrame,
    QFileDialog, QSizePolicy, QTextBrowser,
    QButtonGroup, QRadioButton,
    QDoubleSpinBox, QLineEdit,
)
from PySide6.QtCore import Qt, QThread, Signal, QSettings
from PySide6.QtGui import QPalette

from base_tool import BaseTool
from utils import extract_questions
from theme import Theme
from widgets import AppButton, DropZone, AnimatedButton, AnimatedProgressBar


def _normalize_text(text):
    """对题目文本做统一的规范化处理，降低标点和空白差异带来的误判。"""
    if not text:
        return ""
    text = text.replace("“", '"').replace("”", '"')
    text = text.replace("‘", "'").replace("’", "'")
    text = text.replace("（", "(").replace("）", ")")
    text = text.replace("【", "[").replace("】", "]")
    text = text.replace("：", ":").replace("；", ";")
    text = text.replace("，", ",").replace("。", ".")
    text = text.replace("、", ",")
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"[^\u4e00-\u9fffA-Za-z0-9]+", "", text)
    return text.lower()


def _split_question_parts(q_lines):
    """将题目拆成题干和选项两部分，便于分别计算相似度。"""
    lines = [line.strip() for line in q_lines if str(line).strip()]
    if not lines:
        return "", ""

    stem_lines = []
    option_lines = []
    for line in lines:
        if re.match(r"^([A-Z]|[\u4e00-\u9fff])([.、)）])", line):
            option_lines.append(line)
        else:
            stem_lines.append(line)

    if not stem_lines and lines:
        stem_lines = [lines[0]]
    stem = "\n".join(stem_lines)
    options = "\n".join(option_lines)
    return stem, options


def _char_bigram_overlap(a, b):
    """用字符级 bigram 的 Jaccard 相似度做轻量级语义近似。"""
    if not a or not b:
        return 0.0
    a_grams = {a[i:i + 2] for i in range(len(a) - 1)}
    b_grams = {b[i:i + 2] for i in range(len(b) - 1)}
    if not a_grams and not b_grams:
        return 0.0
    return len(a_grams & b_grams) / len(a_grams | b_grams)


def _token_overlap(a, b):
    """基于字符 token 的重叠率做辅助打分。"""
    if not a or not b:
        return 0.0
    a_tokens = Counter(_normalize_text(a))
    b_tokens = Counter(_normalize_text(b))
    if not a_tokens and not b_tokens:
        return 0.0
    union = set(a_tokens) | set(b_tokens)
    if not union:
        return 0.0
    inter = set(a_tokens) & set(b_tokens)
    return len(inter) / len(union)


def score_question_pair(q_lines_a, q_lines_b):
    """为一对题目计算相似度分数，并返回可解释的结果。"""
    stem_a, options_a = _split_question_parts(q_lines_a)
    stem_b, options_b = _split_question_parts(q_lines_b)

    stem_a_norm = _normalize_text(stem_a)
    stem_b_norm = _normalize_text(stem_b)
    option_a_norm = _normalize_text(options_a)
    option_b_norm = _normalize_text(options_b)
    full_a_norm = _normalize_text("\n".join(q_lines_a))
    full_b_norm = _normalize_text("\n".join(q_lines_b))

    stem_ratio = SequenceMatcher(None, stem_a_norm, stem_b_norm).ratio()
    option_ratio = SequenceMatcher(None, option_a_norm, option_b_norm).ratio()
    full_ratio = SequenceMatcher(None, full_a_norm, full_b_norm).ratio()
    token_ratio = _token_overlap(stem_a_norm, stem_b_norm)
    bigram_ratio = _char_bigram_overlap(stem_a_norm, stem_b_norm)

    score = 0.5 * full_ratio + 0.25 * stem_ratio + 0.15 * token_ratio + 0.1 * bigram_ratio

    if option_a_norm and option_b_norm:
        score = max(score, 0.55 * stem_ratio + 0.45 * option_ratio)

    if option_a_norm and option_b_norm and option_ratio >= 0.9 and stem_ratio >= 0.5:
        score += 0.15
    elif stem_ratio >= 0.7 and full_ratio >= 0.75:
        score += 0.05

    if stem_ratio >= 0.8 and (option_ratio >= 0.8 or bigram_ratio >= 0.7):
        score = max(score, 0.86)
    if stem_ratio >= 0.95 and option_ratio >= 0.9:
        score = max(score, 0.95)

    score = min(1.0, max(0.0, score))

    reason = "低相似"
    if score >= 0.9:
        reason = "高度相似"
    elif score >= 0.8:
        reason = "较高相似"
    elif score >= 0.7:
        reason = "中等相似"

    return {
        "score": round(score, 4),
        "reason": reason,
        "stem_ratio": round(stem_ratio, 4),
        "option_ratio": round(option_ratio, 4),
        "full_ratio": round(full_ratio, 4),
        "token_ratio": round(token_ratio, 4),
        "bigram_ratio": round(bigram_ratio, 4),
    }


# ── multi-file drop zone ───────────────────────────────────────────

class MultiDropZone(DropZone):
    """扩展 DropZone，支持同时选择和拖入多个文件，显示已选文件数量。"""

    files_selected = Signal(list)

    def __init__(self, placeholder_text, file_filter="", compact=False, theme=None):
        super().__init__(placeholder_text, file_filter, compact, theme)

    def _open_dialog(self):
        paths, _ = QFileDialog.getOpenFileNames(
            self, "选择文件", "", self._file_filter,
        )
        if paths:
            self._handle_files(paths)

    def dropEvent(self, event):
        self.setStyleSheet(self._normal_style)
        urls = event.mimeData().urls()
        if urls:
            paths = [url.toLocalFile() for url in urls]
            self._handle_files(paths)

    def _handle_files(self, paths):
        self._paths = paths
        if len(paths) == 1:
            self.set_file(paths[0])
        else:
            self.label.setText(f"已选择 {len(paths)} 个文件")
            self._apply_style()
        self.files_selected.emit(paths)


# ── worker ─────────────────────────────────────────────────────────

class CheckerWorker(QThread):
    """后台查重线程，逐个比对主文档与副文档中的题目。"""

    log = Signal(str)
    finished = Signal(dict)

    def __init__(self, main_path, secondary_paths, num_pat, opt_pre, threshold):
        super().__init__()
        self.main_path = main_path
        self.secondary_paths = secondary_paths
        self.num_pat = num_pat
        self.opt_pre = opt_pre
        self.threshold = threshold

    def run(self):
        try:
            self.log.emit("正在提取主文档题目…")
            main_qs = extract_questions(self.main_path, self.num_pat, self.opt_pre)
            if not main_qs:
                self.finished.emit({"error": "主文档中未提取到题目"})
                return

            self.log.emit(f"主文档：共 {len(main_qs)} 道题")
            self.log.emit(f"模糊匹配阈值：{self.threshold:.0%}")

            secondary_data = {}
            for path in self.secondary_paths:
                fname = os.path.basename(path)
                self.log.emit(f"正在提取副文档：{fname}")
                qs = extract_questions(path, self.num_pat, self.opt_pre)
                if qs:
                    # 保留题目原始行列表（List[str]），与 score_question_pair 的契约一致。
                    # 注意：不要在此做字符串化预处理，否则会破坏选项相似度分支，
                    # 并导致与多对多模式（同样传入 List[str]）结果不一致。
                    secondary_data[fname] = [list(q) for q in qs]
                    self.log.emit(f"  {fname}：{len(qs)} 道题")
                else:
                    self.log.emit(f"  {fname}：未提取到题目")

            if not secondary_data:
                self.finished.emit({"error": "所有副文档均未提取到题目"})
                return

            total = len(main_qs)
            details = []
            for idx, q_lines in enumerate(main_qs, 1):
                if idx % 5 == 0:
                    self.log.emit(f"比对进度：{idx}/{total}")
                sources = []
                for fname, sec_texts in secondary_data.items():
                    matched = False
                    for s_text in sec_texts:
                        result = score_question_pair(q_lines, s_text)
                        if result["score"] >= self.threshold:
                            sources.append({
                                "file": fname,
                                "score": result["score"],
                                "reason": result["reason"],
                            })
                            matched = True
                            break
                if sources:
                    details.append({
                        "index": idx,
                        "text": q_lines,
                        "sources": sources,
                    })

            self.log.emit(f"检测完成：{total} 道题目中，重复 {len(details)} 道")
            self.finished.emit({
                "mode": "1_to_many",
                "main_count": total,
                "duplicate_count": len(details),
                "details": details,
            })
        except Exception as e:
            self.finished.emit({"error": str(e)})


# ── many-to-many worker ────────────────────────────────────────────

class ManyToManyWorker(QThread):
    """多对多查重线程：对所有文档内的题目做两两比对，检测文档内及跨文档重复。"""

    log = Signal(str)
    finished = Signal(dict)

    def __init__(self, all_paths, num_pat, opt_pre, threshold):
        super().__init__()
        self.all_paths = all_paths
        self.num_pat = num_pat
        self.opt_pre = opt_pre
        self.threshold = threshold

    def run(self):
        try:
            n_docs = len(self.all_paths)
            self.log.emit(f"多对多模式：正在分析 {n_docs} 份文档…")

            all_questions = []
            doc_questions = {}

            for doc_idx, path in enumerate(self.all_paths):
                fname = os.path.basename(path)
                self.log.emit(f"  正在提取：{fname}")
                qs = extract_questions(path, self.num_pat, self.opt_pre)
                doc_questions[fname] = len(qs)
                self.log.emit(f"    {fname}：{len(qs)} 道题")
                for q_idx, q_lines in enumerate(qs):
                    all_questions.append((fname, doc_idx, q_idx, q_lines))

            total_questions = len(all_questions)
            if total_questions < 2:
                self.finished.emit({"error": "所有文档题目总数不足 2 道，无法比对"})
                return

            self.log.emit(f"总题目数：{total_questions}，开始两两比对…")

            duplicate_pairs = []
            total_pairs = total_questions * (total_questions - 1) // 2
            checked = 0

            for i in range(total_questions):
                fname_i, doc_i, qi, q_lines_i = all_questions[i]
                for j in range(i + 1, total_questions):
                    fname_j, doc_j, qj, q_lines_j = all_questions[j]
                    checked += 1
                    if checked % 500 == 0:
                        self.log.emit(f"  比对进度：{checked}/{total_pairs}")

                    result = score_question_pair(q_lines_i, q_lines_j)
                    if result["score"] >= self.threshold:
                        pair_type = "internal" if doc_i == doc_j else "cross"
                        duplicate_pairs.append({
                            "q1": {"file": fname_i, "index": qi + 1, "text": q_lines_i},
                            "q2": {"file": fname_j, "index": qj + 1, "text": q_lines_j},
                            "score": result["score"],
                            "reason": result["reason"],
                            "type": pair_type,
                        })

            self.log.emit(f"比对完成：{total_pairs} 对中，发现 {len(duplicate_pairs)} 对重复")

            self.finished.emit({
                "mode": "many_to_many",
                "total_questions": total_questions,
                "document_count": n_docs,
                "doc_questions": doc_questions,
                "duplicate_pairs": duplicate_pairs,
                "duplicate_rate": round(len(duplicate_pairs) / max(total_pairs, 1), 4),
            })

        except Exception as e:
            self.finished.emit({"error": str(e)})


# ── tool ────────────────────────────────────────────────────────────

class SimilarityCheckerTool(BaseTool):
    def __init__(self):
        super().__init__()
        self.theme = Theme()
        self._main_path = ""
        self._secondary_paths = []
        self._all_paths = []
        self._mode = "1_to_many"
        self._worker = None
        # 查重参数持久化（阈值 / 题号格式 / 选项前缀），提供默认值与重置
        self.settings = QSettings("SimilarityChecker", "SimilarityChecker")
        self._setup_background()
        self._setup_ui()
        QApplication.instance().styleHints().colorSchemeChanged.connect(
            self._on_theme_changed
        )

    # ── interface ───────────────────────────────────────────────

    def get_name(self) -> str:
        return "Similarity Checker"

    def get_description(self) -> str:
        return "检测主文档与多个副文档之间的题目重复率，支持精确/模糊匹配，导出查重报告。"

    # ── theme ───────────────────────────────────────────────────

    def _setup_background(self):
        pal = self.palette()
        pal.setColor(QPalette.Window, self.theme.window_solid_bg)
        self.setPalette(pal)
        self.setAutoFillBackground(True)

    def _on_theme_changed(self):
        self.theme.refresh()
        self._restyle_all()

    def _restyle_all(self):
        t = self.theme
        pal = self.palette()
        pal.setColor(QPalette.Window, t.window_solid_bg)
        self.setPalette(pal)

        self._main_card.setStyleSheet(t.qss_card())
        for dz in [self._main_drop_zone, self._secondary_drop_zone, self._all_drop_zone]:
            dz._theme = t
            dz._apply_style()

        radio_style = (
            f"QRadioButton {{ color: {t.text_primary}; font-size: 14px; "
            f"background: transparent; spacing: 6px; padding: 4px 12px; }}"
            f"QRadioButton::indicator {{ width: 16px; height: 16px; }}"
        )
        for rb in [self._radio_1toN, self._radio_NtoN]:
            rb.setStyleSheet(radio_style)

        header_s = t.qss_section_header()
        for lbl in self._section_labels:
            lbl.setStyleSheet(header_s)

        self._divider.setStyleSheet(t.qss_divider())

        self._check_btn.set_theme(t)
        self._export_btn.set_theme(t)

        self._progress_bar.setStyleSheet(t.qss_progress_bar())

        self._log_browser.setStyleSheet(
            f"QTextBrowser {{ background: {t.input_bg}; color: {t.text_primary}; "
            f"border: none; border-radius: 8px; padding: 12px; font-size: 13px; }}"
        )

        input_s = (
            f"QLineEdit, QDoubleSpinBox {{ padding: 4px 8px; border: none; "
            f"border-radius: 8px; font-size: 14px; background: {t.input_bg}; "
            f"color: {t.text_primary}; }}"
        )
        self._threshold_spin.setStyleSheet(input_s)
        self._num_edit.setStyleSheet(input_s)
        self._opt_edit.setStyleSheet(input_s)
        self._reset_btn.setStyleSheet(
            f"QPushButton {{ padding: 4px 12px; border: none; border-radius: 8px; "
            f"font-size: 14px; background: {t.input_bg}; color: {t.text_primary}; }}"
        )

        self._export_btn.set_theme(t)

    # ── ui ──────────────────────────────────────────────────────

    def _setup_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(16, 16, 16, 16)
        root.setSpacing(0)

        self._section_labels = []

        card = QFrame()
        self._main_card = card
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(24, 20, 24, 20)
        card_layout.setSpacing(16)

        t = self.theme

        # ── 模式选择 ──
        mode_row = QHBoxLayout()
        mode_row.setSpacing(4)
        mode_label = QLabel("查重模式：")
        mode_label.setStyleSheet(
            f"font-size: 14px; font-weight: bold; color: {t.card_header_color}; "
            "background: transparent; padding: 0;"
        )
        mode_row.addWidget(mode_label)

        self._mode_group = QButtonGroup(self)
        self._radio_1toN = QRadioButton("1对多")
        self._radio_NtoN = QRadioButton("多对多")
        self._radio_1toN.setChecked(True)
        self._mode_group.addButton(self._radio_1toN, 0)
        self._mode_group.addButton(self._radio_NtoN, 1)
        self._mode_group.buttonClicked.connect(self._on_mode_changed)

        mode_row.addWidget(self._radio_1toN)
        mode_row.addWidget(self._radio_NtoN)
        mode_row.addStretch()
        card_layout.addLayout(mode_row)

        # ── 1对多模式 UI ──
        self._one_to_many_widgets = []

        lbl_main = self._section_label("📄 主文档（选择题库）", card_layout)
        self._one_to_many_widgets.append(lbl_main)
        self._main_drop_zone = DropZone("点击或拖拽 .docx 文件", "Word 文档 (*.docx)", theme=t)
        self._main_drop_zone.file_selected.connect(self._on_main_file)
        self._one_to_many_widgets.append(self._main_drop_zone)
        card_layout.addWidget(self._main_drop_zone)

        self._divider = QFrame()
        self._divider.setFixedHeight(1)
        self._divider.setStyleSheet(f"background: {t.border}; border: none;")
        self._one_to_many_widgets.append(self._divider)
        card_layout.addWidget(self._divider)

        lbl_sec = self._section_label("📑 副文档（对比库，可多选）", card_layout)
        self._one_to_many_widgets.append(lbl_sec)
        self._secondary_drop_zone = MultiDropZone("点击或拖拽 .docx 文件（可多选）", "Word 文档 (*.docx)", theme=t)
        self._secondary_drop_zone.files_selected.connect(self._on_secondary_files)
        self._one_to_many_widgets.append(self._secondary_drop_zone)
        card_layout.addWidget(self._secondary_drop_zone)

        # ── 多对多模式 UI ──
        self._many_to_many_widgets = []

        lbl_all = self._section_label("📚 所有文档（可多选，至少2份）", card_layout)
        self._many_to_many_widgets.append(lbl_all)
        self._all_drop_zone = MultiDropZone("点击或拖拽 .docx 文件（可多选）", "Word 文档 (*.docx)", theme=t)
        self._all_drop_zone.files_selected.connect(self._on_all_files)
        self._many_to_many_widgets.append(self._all_drop_zone)
        card_layout.addWidget(self._all_drop_zone)

        # 默认隐藏多对多 UI
        for w in self._many_to_many_widgets:
            w.setVisible(False)

        # ── 高级设置（阈值 / 题号格式 / 选项前缀，可持久化 + 重置） ──
        settings_row = QHBoxLayout()
        settings_row.setSpacing(8)

        th_label = QLabel("相似度阈值")
        self._threshold_spin = QDoubleSpinBox()
        self._threshold_spin.setRange(0.5, 1.0)
        self._threshold_spin.setSingleStep(0.01)
        self._threshold_spin.setDecimals(2)
        self._threshold_spin.setValue(float(self.settings.value("threshold", 0.8)))

        num_label = QLabel("题号格式")
        self._num_edit = QLineEdit()
        self._num_edit.setPlaceholderText("如 1.")
        self._num_edit.setText(self.settings.value("num_pattern", "1."))

        opt_label = QLabel("选项前缀")
        self._opt_edit = QLineEdit()
        self._opt_edit.setPlaceholderText("如 A.")
        self._opt_edit.setText(self.settings.value("opt_prefix", "A."))

        self._reset_btn = QPushButton("重置")
        self._reset_btn.clicked.connect(self._on_reset_settings)

        settings_row.addWidget(th_label)
        settings_row.addWidget(self._threshold_spin)
        settings_row.addWidget(num_label)
        settings_row.addWidget(self._num_edit)
        settings_row.addWidget(opt_label)
        settings_row.addWidget(self._opt_edit)
        settings_row.addWidget(self._reset_btn)
        settings_row.addStretch()
        card_layout.addLayout(settings_row)

        # 参数变更即持久化（默认值 0.8 / "1." / "A."）
        self._threshold_spin.valueChanged.connect(self._save_settings)
        self._num_edit.textChanged.connect(self._save_settings)
        self._opt_edit.textChanged.connect(self._save_settings)

        # ── 操作按钮 ──
        btn_row = QHBoxLayout()
        btn_row.setSpacing(12)
        self._check_btn = AnimatedButton("开始检测", default_height=44, theme=self.theme)
        self._check_btn.clicked.connect(self._on_check)
        self._export_btn = AppButton("导出报告", default_height=44, theme=self.theme)
        self._export_btn.clicked.connect(self._on_export)
        self._export_btn.set_actionable(False, "请先完成检测后导出")
        btn_row.addWidget(self._check_btn)
        btn_row.addWidget(self._export_btn)
        card_layout.addLayout(btn_row)

        card_layout.addSpacing(8)

        self._progress_bar = AnimatedProgressBar()
        self._progress_bar.setTextVisible(False)
        self._progress_bar.setVisible(False)
        card_layout.addWidget(self._progress_bar)

        self._log_browser = QTextBrowser()
        self._log_browser.setOpenExternalLinks(False)
        self._log_browser.setMinimumHeight(180)
        self._log_browser.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        card_layout.addWidget(self._log_browser)

        root.addWidget(card)
        self._restyle_all()

    def _section_label(self, text, layout):
        label = QLabel(text)
        label.setStyleSheet(
            f"font-size: 15px; font-weight: bold; color: {self.theme.card_header_color}; "
            "background: transparent; padding: 0;"
        )
        self._section_labels.append(label)
        layout.addWidget(label)
        return label

    # ── events ──────────────────────────────────────────────────

    def _on_mode_changed(self, button):
        if button == self._radio_1toN:
            self._mode = "1_to_many"
            for w in self._one_to_many_widgets:
                w.setVisible(True)
            for w in self._many_to_many_widgets:
                w.setVisible(False)
        else:
            self._mode = "many_to_many"
            for w in self._one_to_many_widgets:
                w.setVisible(False)
            for w in self._many_to_many_widgets:
                w.setVisible(True)
        self._log_browser.clear()

    def _on_main_file(self, path):
        self._main_path = path
        self._log_browser.clear()

    def _on_secondary_files(self, paths):
        self._secondary_paths = paths
        self._log_browser.clear()

    def _on_all_files(self, paths):
        self._all_paths = paths
        self._log_browser.clear()

    def _save_settings(self):
        """将查重参数持久化到 QSettings。"""
        self.settings.setValue("threshold", self._threshold_spin.value())
        self.settings.setValue("num_pattern", self._num_edit.text())
        self.settings.setValue("opt_prefix", self._opt_edit.text())

    def _on_reset_settings(self):
        """重置查重参数为默认值（阈值 0.8 / 题号 "1." / 选项 "A."）。"""
        self._threshold_spin.setValue(0.8)
        self._num_edit.setText("1.")
        self._opt_edit.setText("A.")
        self._save_settings()

    def _on_check(self):
        self._log_browser.clear()
        self._check_btn.set_loading(True)
        self._export_btn.setEnabled(False)
        self._progress_bar.setVisible(True)
        self._progress_bar.setRange(0, 0)
        self._progress_bar.setValue(0)

        if self._mode == "1_to_many":
            self._start_one_to_many()
        else:
            self._start_many_to_many()

    def _start_one_to_many(self):
        self._stop_worker()
        if not self._main_path:
            self._log_browser.setHtml(
                f"<p style='color:{self.theme.error_color};'>请先选择主文档</p>"
            )
            self._check_btn.set_loading(False)
            self._progress_bar.setVisible(False)
            return
        if not self._secondary_paths:
            self._log_browser.setHtml(
                f"<p style='color:{self.theme.error_color};'>请先选择至少一个副文档</p>"
            )
            self._check_btn.set_loading(False)
            self._progress_bar.setVisible(False)
            return

        self._worker = CheckerWorker(
            self._main_path,
            self._secondary_paths,
            self._num_edit.text().strip() or "1.",
            self._opt_edit.text().strip() or "A.",
            self._threshold_spin.value(),
        )
        self._worker.log.connect(self._on_log)
        self._worker.finished.connect(self._on_finished)
        self._worker.start()

    def _start_many_to_many(self):
        self._stop_worker()
        if len(self._all_paths) < 2:
            self._log_browser.setHtml(
                f"<p style='color:{self.theme.error_color};'>请先选择至少 2 份文档</p>"
            )
            self._check_btn.set_loading(False)
            self._progress_bar.setVisible(False)
            return

        self._worker = ManyToManyWorker(
            self._all_paths,
            self._num_edit.text().strip() or "1.",
            self._opt_edit.text().strip() or "A.",
            self._threshold_spin.value(),
        )
        self._worker.log.connect(self._on_log)
        self._worker.finished.connect(self._on_finished)
        self._worker.start()

    def _on_log(self, msg):
        self._log_browser.append(msg)

    def _on_finished(self, result):
        # 忽略已停止/废弃的旧线程回调，避免重复结果污染状态
        if self.sender() is not self._worker:
            return
        self._check_btn.set_loading(False)
        self._progress_bar.setVisible(False)
        self._export_btn.setEnabled(False)

        if "error" in result:
            self._log_browser.append(f"\n错误：{result['error']}")
            return

        mode = result.get("mode", "1_to_many")

        if mode == "many_to_many":
            self._display_many_to_many_result(result)
        else:
            self._display_one_to_many_result(result)

        self._last_result = result
        self._last_mode = mode
        if (mode == "1_to_many" and result.get("duplicate_count", 0) > 0) or \
           (mode == "many_to_many" and len(result.get("duplicate_pairs", [])) > 0):
            self._export_btn.setEnabled(True)

        # 释放后台线程对象，避免长期持有造成泄漏
        if self._worker is not None:
            self._worker.deleteLater()
            self._worker = None

    def stop_worker(self):
        """供主窗口 closeEvent 调用，统一清理后台查重线程。"""
        self._stop_worker()

    def _stop_worker(self):
        """停止并清理正在运行的后台线程，避免重复点击产生孤儿线程。

        CPU 密集线程没有事件循环，quit() 无法中断其 run()，故若仍在运行则
        wait() 等待其自然结束（窗口关闭场景下保证线程被回收，不沦为孤儿），
        随后 deleteLater() 释放对象。先 disconnect 旧信号，避免旧线程的
        残留回调污染 UI 状态。
        """
        w = getattr(self, '_worker', None)
        if w is None:
            return
        try:
            w.log.disconnect()
            w.finished.disconnect()
        except (TypeError, RuntimeError):
            pass
        w.quit()
        if w.isRunning():
            w.wait()
        w.deleteLater()
        self._worker = None

    def _display_one_to_many_result(self, result):
        main_count = result["main_count"]
        dup_count = result["duplicate_count"]
        details = result["details"]

        self._log_browser.append("")
        self._log_browser.append(f"──── 检测摘要 ────")
        self._log_browser.append(f"主文档题目数：{main_count}")
        self._log_browser.append(f"重复题目数：{dup_count}")
        self._log_browser.append(f"重复率：{dup_count / max(main_count, 1) * 100:.1f}%")
        self._log_browser.append("")

        for d in details:
            text_preview = d["text"][0][:50] + "…" if len(d["text"][0]) > 50 else d["text"][0]
            self._log_browser.append(
                f"第{d['index']}题 - {text_preview}"
            )
            for item in d["sources"]:
                self._log_browser.append(
                    f"  重复来源：{item['file']} (相似度 {item['score']:.2f}, {item['reason']})"
                )
            self._log_browser.append("")

    def _display_many_to_many_result(self, result):
        total_qs = result["total_questions"]
        n_docs = result["document_count"]
        doc_qs = result["doc_questions"]
        pairs = result["duplicate_pairs"]
        rate = result["duplicate_rate"]

        internal_count = sum(1 for p in pairs if p["type"] == "internal")
        cross_count = sum(1 for p in pairs if p["type"] == "cross")

        self._log_browser.append("")
        self._log_browser.append(f"──── 检测摘要 ────")
        self._log_browser.append(f"文档数：{n_docs}，总题目数：{total_qs}")
        self._log_browser.append(f"重复对总数：{len(pairs)}（文档内 {internal_count}，跨文档 {cross_count}）")
        self._log_browser.append("")
        self._log_browser.append("文档题目分布：")
        for fname, count in doc_qs.items():
            self._log_browser.append(f"  {fname}：{count} 题")
        self._log_browser.append("")

        for i, pair in enumerate(pairs, 1):
            tag = "[跨文档]" if pair["type"] == "cross" else "[文档内]"
            q1_text = pair["q1"]["text"][0][:40] + "…" if len(pair["q1"]["text"][0]) > 40 else pair["q1"]["text"][0]
            self._log_browser.append(
                f"{i}. {tag} {pair['q1']['file']}-第{pair['q1']['index']}题 "
                f"⇄ {pair['q2']['file']}-第{pair['q2']['index']}题 "
                f"({pair['score']:.2f}, {pair['reason']})"
            )
            self._log_browser.append(f"   {q1_text}")
            self._log_browser.append("")

    def _on_export(self):
        if not hasattr(self, '_last_result') or not self._last_result:
            return

        mode = getattr(self, '_last_mode', '1_to_many')
        result = self._last_result

        if mode == "many_to_many":
            self._export_many_to_many(result)
        else:
            self._export_one_to_many(result)

    def _export_one_to_many(self, result):
        start_dir = os.path.dirname(self._main_path) if self._main_path else ""
        default_name = os.path.join(start_dir, "查重报告.docx")
        path, _ = QFileDialog.getSaveFileName(
            self, "导出查重报告", default_name, "Word 文档 (*.docx)",
        )
        if not path:
            return

        doc = Document()

        doc.add_heading("题目查重报告（1对多模式）", 0)
        doc.add_paragraph(
            f"主文档题目数：{result['main_count']}，"
            f"重复题目数：{result['duplicate_count']}，"
            f"重复率：{result['duplicate_count'] / max(result['main_count'], 1) * 100:.1f}%"
        )

        for d in result["details"]:
            doc.add_heading(f"第 {d['index']} 题", 2)
            for line in d["text"]:
                doc.add_paragraph(line, style="List Bullet")
            doc.add_paragraph(
                "重复来源：" + "; ".join(
                    f"{item['file']} ({item['score']:.2f}, {item['reason']})"
                    for item in d["sources"]
                )
            )

        doc.save(path)
        self._open_folder(path)

    def _export_many_to_many(self, result):
        start_dir = os.path.dirname(self._all_paths[0]) if self._all_paths else ""
        default_name = os.path.join(start_dir, "查重报告.docx")
        path, _ = QFileDialog.getSaveFileName(
            self, "导出查重报告", default_name, "Word 文档 (*.docx)",
        )
        if not path:
            return

        doc = Document()
        pairs = result["duplicate_pairs"]
        internal_count = sum(1 for p in pairs if p["type"] == "internal")
        cross_count = sum(1 for p in pairs if p["type"] == "cross")

        doc.add_heading("题目查重报告（多对多模式）", 0)
        doc.add_paragraph(
            f"文档数：{result['document_count']}，"
            f"总题目数：{result['total_questions']}，"
            f"重复对总数：{len(pairs)}（文档内 {internal_count}，跨文档 {cross_count}）"
        )

        doc.add_heading("文档题目分布", 2)
        for fname, count in result["doc_questions"].items():
            doc.add_paragraph(f"{fname}：{count} 题", style="List Bullet")

        doc.add_heading("重复详情", 2)
        for i, pair in enumerate(pairs, 1):
            tag = "跨文档" if pair["type"] == "cross" else "文档内"
            doc.add_heading(
                f"{i}. [{tag}] {pair['q1']['file']}-第{pair['q1']['index']}题 "
                f"⇄ {pair['q2']['file']}-第{pair['q2']['index']}题 "
                f"({pair['score']:.2f}, {pair['reason']})",
                3,
            )
            doc.add_paragraph(
                f"题目 1（{pair['q1']['file']} 第{pair['q1']['index']}题）："
            )
            for line in pair["q1"]["text"]:
                doc.add_paragraph(line, style="List Bullet")
            doc.add_paragraph(
                f"题目 2（{pair['q2']['file']} 第{pair['q2']['index']}题）："
            )
            for line in pair["q2"]["text"]:
                doc.add_paragraph(line, style="List Bullet")
            doc.add_paragraph("")

        doc.save(path)
        self._open_folder(path)

    def _open_folder(self, path):
        folder = os.path.dirname(path) or "."
        try:
            if sys.platform == "darwin":
                subprocess.Popen(["open", folder])
            elif sys.platform == "win32":
                subprocess.Popen(["explorer", folder])
            else:
                subprocess.Popen(["xdg-open", folder])
        except Exception:
            pass

    def showEvent(self, event):
        super().showEvent(event)
