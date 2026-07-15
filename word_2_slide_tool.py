import html as _html
import logging
import os
import re as _re
import subprocess
import sys
import tempfile

from docx import Document
from pptx import Presentation
from PySide6.QtWidgets import (
    QApplication, QWidget, QLineEdit,
    QPushButton, QLabel, QVBoxLayout, QHBoxLayout,
    QFileDialog, QSizePolicy, QDialog,
    QFrame, QComboBox, QSpinBox,
    QGraphicsDropShadowEffect, QDoubleSpinBox, QCheckBox,
    QScrollArea, QDialogButtonBox, QTextBrowser
)
from PySide6.QtCore import (
    Qt, QThread, Signal, QSettings,
    QPropertyAnimation, QEasingCurve,
)
from PySide6.QtGui import (
    QPalette, QIcon,
)

from base_tool import BaseTool
from theme import Theme, _get_system_fonts
from utils import extract_questions, _resolve_line_spacing, generate_pptx
from widgets import AppButton, AnimatedButton, AnimatedProgressBar, ToastNotification, DropZone

logger = logging.getLogger(__name__)


# ── 预览 HTML 安全转义 ─────────────────────────────────────────────
# QTextBrowser.setHtml 会解析富文本。用户文档中的 <script> / <img src onerror=...>
# 等标签可能触发资源加载或布局破坏；字体名进入 <style> 的 CSS 字符串上下文，
# 仅用 html.escape 无法防御 { } 注入。两个辅助函数分别处理正文与 CSS 上下文。

# 仅保留的白名单格式化标签：<b> <i> <u> <br>（含 </b> 与 <br/> 变体）
_PREVIEW_SAFE_TAG_RE = _re.compile(r"&lt;(/?)(br|b|i|u)\b\s*(/?)&gt;", _re.IGNORECASE)


def _escape_preview_line(text: str) -> str:
    """转义题面文本用于预览，仅保留 <b>/<i>/<u>/<br> 白名单标签。

    html.escape 把 < > & ' " 全部实体化，先杜绝任何标签/属性注入；
    再用白名单正则把安全的格式化标签还原回来。
    """
    escaped = _html.escape(text)
    return _PREVIEW_SAFE_TAG_RE.sub(r"<\1\2\3>", escaped)


def _sanitize_font_name(font_name: str) -> str:
    """净化字体名（进入 <style> CSS 字符串上下文），防御 CSS 注入。

    字体名来自可编辑下拉框（用户可任意输入），先剔除能脱离 CSS 字符串或
    开启新规则的字符 { } ' " ` ;，再 html.escape 处理 < > & 等。
    """
    stripped = _re.sub(r"[{}\"'`;]", "", font_name or "")
    return _html.escape(stripped)


# ── worker ─────────────────────────────────────────────────────────

class ConvertWorker(QThread):
    progress_text = Signal(str)
    finished = Signal(bool, str)

    def __init__(self, ppt_path, questions, font_name, font_size, out_path,
                 line_spacing_type, line_spacing_value, first_line_indent):
        super().__init__()
        self.ppt_path = ppt_path
        self.questions = questions
        self.font_name = font_name
        self.font_size = font_size
        self.out_path = out_path
        self.line_spacing_type = line_spacing_type
        self.line_spacing_value = line_spacing_value
        self.first_line_indent = first_line_indent

    def run(self):
        """线程执行入口，调用 generate_pptx 生成 PPT 文件。

        成功时发射 finished(True, 消息)，异常时发射 finished(False, 错误信息)。
        处理过程中通过 progress_text 信号实时报告进度。
        """
        try:
            generate_pptx(
                self.ppt_path, self.questions,
                self.font_name, self.font_size, self.out_path,
                line_spacing_type=self.line_spacing_type,
                line_spacing_value=self.line_spacing_value,
                first_line_indent=self.first_line_indent,
                progress_cb=lambda cur, total: self.progress_text.emit(
                    f"正在处理第 {cur}/{total} 道题"
                ),
            )
            self.finished.emit(True, f"生成成功，共 {len(self.questions) * 2} 页")
        except Exception as e:
            self.finished.emit(False, str(e))


class ExtractWorker(QThread):
    """在后台线程中解析 Word 文档、提取题目，避免大文档在 GUI 线程同步解析导致卡顿。

    成功时发射 extracted(questions)，异常时发射 error(message)。
    CPU 密集线程无事件循环，quit() 无法中断其 run()，故停止时直接 wait() 回收。
    """

    extracted = Signal(list)
    error = Signal(str)

    def __init__(self, doc_path, num_pat, opt_pre):
        super().__init__()
        self.doc_path = doc_path
        self.num_pat = num_pat
        self.opt_pre = opt_pre

    def run(self):
        try:
            qs = extract_questions(self.doc_path, self.num_pat, self.opt_pre)
            self.extracted.emit(qs)
        except Exception as e:
            self.error.emit(str(e))


# ── tool ────────────────────────────────────────────────────────────

class Quiz2SlideTool(BaseTool):
    def get_name(self):
        return "Quiz2Slide"

    def get_description(self):
        return "将 Word 题目文档转换为可直接使用的 PowerPoint 幻灯片。"

    def __init__(self):
        """初始化主窗口，创建主题管理器、UI 组件并监听系统配色变化。"""
        super().__init__()
        self.theme = Theme()
        self.setWindowTitle("Quiz2Slide")
        self.resize(700, 700)
        self.setMinimumSize(680, 600)
        self.settings = QSettings("Quiz2Slide", "Quiz2Slide")
        self._word_path = ""
        self._ppt_path = ""
        self.worker = None
        self._extract_worker = None
        self._setup_background()
        self._setup_ui()
        self._center_on_screen()
        QApplication.instance().styleHints().colorSchemeChanged.connect(
            self._on_theme_changed
        )

    # ── background ──────────────────────────────────────────────

    def _setup_background(self):
        """设置窗口初始背景色（使用主题色）。"""
        pal = self.palette()
        pal.setColor(QPalette.Window, self.theme.window_solid_bg)
        self.setPalette(pal)
        self.setAutoFillBackground(True)

    def _center_on_screen(self):
        """将窗口移动到屏幕中央。"""
        screen = QApplication.primaryScreen().availableGeometry()
        x = (screen.width() - self.width()) // 2
        y = (screen.height() - self.height()) // 2
        self.move(x, y)

    def _on_theme_changed(self):
        """系统配色方案变化时的回调，刷新主题并重绘所有组件。"""
        self.theme.refresh()
        self._restyle_all()

    def _input_base_style(self, t, widget_type, extra=""):
        """生成输入类控件（QLineEdit/QSpinBox/QComboBox）的基础样式表。

        Args:
            t: Theme 对象。
            widget_type: 控件类型字符串，如 "QLineEdit"。
            extra: 附加在该控件基础样式之后的子控件样式字符串。

        Returns:
            完整的 QSS 样式表字符串。
        """
        base = (
            f"{widget_type} {{ padding: 4px 8px; border: none; "
            f"border-radius: 8px; font-size: 14px; background: {t.input_bg}; "
            f"color: {t.text_primary}; }}"
            f"{widget_type}:focus {{ border: 1px solid {t.accent}; "
            f"background: {t.card_bg}; }}"
        )
        return base + extra

    def _restyle_all(self):
        """根据当前主题重新应用所有组件的样式表。

        包括窗口背景、输入框、下拉框、卡片、标签、按钮、进度条、
        拖放区域等所有 UI 元素的样式更新。
        """
        t = self.theme
        pal = self.palette()
        pal.setColor(QPalette.Window, t.window_solid_bg)
        self.setPalette(pal)

        input_s = self._input_base_style(t, "QLineEdit")
        spin_s = self._input_base_style(
            t, "QSpinBox",
            f"QSpinBox::up-button {{ width: 22px; height: 12px; "
            f"subcontrol-position: top right; "
            f"border: none; background: transparent; }}"
            f"QSpinBox::down-button {{ width: 22px; height: 12px; "
            f"subcontrol-position: bottom right; "
            f"border: none; background: transparent; }}"
        )
        combo_s = self._input_base_style(
            t, "QComboBox",
            f"QComboBox::drop-down {{ border: none; width: 24px; }}"
            f"QComboBox QAbstractItemView {{ border: 1px solid {t.border}; "
            f"border-radius: 8px; selection-background-color: {t.accent}; padding: 4px; }}"
        )
        dspin_s = (
            f"QDoubleSpinBox {{ padding: 4px 8px; border: none; "
            f"border-radius: 8px; font-size: 14px; background: {t.input_bg}; "
            f"color: {t.text_primary}; }}"
            f"QDoubleSpinBox:focus {{ border: 1px solid {t.accent}; "
            f"background: {t.card_bg}; }}"
        )
        check_s = f"QCheckBox {{ color: {t.text_primary}; }}"

        for w in [self.question_num_fmt, self.option_prefix]:
            w.setStyleSheet(input_s)
        self.font_size.setStyleSheet(spin_s)
        self.font_name.setStyleSheet(combo_s)
        self.line_spacing_type.setStyleSheet(combo_s)
        self.line_spacing_value.setStyleSheet(dspin_s)
        self.first_line_indent.setStyleSheet(check_s)

        self._main_card.setStyleSheet(t.qss_card())
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(20)
        shadow.setOffset(0, 4)
        shadow.setColor(t.shadow_color)
        self._main_card.setGraphicsEffect(shadow)

        label_s = f"font-size: 12px; color: {t.text_secondary}; margin-bottom: 2px;"
        for lbl in self._field_labels:
            lbl.setStyleSheet(label_s)

        header_s = t.qss_section_header()
        for lbl in self._section_labels:
            lbl.setStyleSheet(header_s)

        save_label_s = (
            f"color: {t.text_secondary}; font-size: 13px; background: transparent;"
        )
        self._save_to_label.setStyleSheet(save_label_s)
        self.out_path_label.setStyleSheet(save_label_s)

        self._change_btn.set_theme(t)

        self._divider.setStyleSheet(t.qss_divider())

        self.error_label.setStyleSheet(f"color: {t.error_color}; font-size: 12px;")
        self.progress_label.setStyleSheet(f"color: {t.text_secondary}; font-size: 12px;")
        self.progress_bar.setStyleSheet(t.qss_progress_bar())


        self.convert_btn.set_theme(t)
        self._change_btn.set_theme(t)

        if hasattr(self, 'word_drop_zone'):
            self.word_drop_zone._apply_style()
        if hasattr(self, 'ppt_drop_zone'):
            self.ppt_drop_zone._apply_style()

        if hasattr(self, '_scroll'):
            self._scroll.setStyleSheet(
                "QScrollArea { background: transparent; border: none; }"
                "QScrollBar:vertical { width: 6px; background: transparent; }"
                f"QScrollBar::handle:vertical {{ background: {t.scrollbar_handle}; "
                "border-radius: 3px; min-height: 30px; }"
                "QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }"
            )

        self.update()

    # ── ui setup ────────────────────────────────────────────────

    def _setup_ui(self):
        """构建整体 UI 布局，包括主内容区和 Toast 通知组件。"""
        self._root = QVBoxLayout(self)
        self._root.setContentsMargins(16, 16, 16, 16)
        self._root.setSpacing(0)
        self._build_main()
        self.toast = ToastNotification(self, theme=self.theme)

    def _build_main(self):
        """构建主卡片内容，包含文件选择、字体设置、行间距、输出路径等所有表单字段。"""
        t = self.theme
        self._field_labels = []
        self._section_labels = []

        card = QFrame()
        self._main_card = card

        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(0, 0, 0, 0)
        card_layout.setSpacing(0)

        self._scroll = QScrollArea()
        self._scroll.setWidgetResizable(True)
        self._scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self._scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self._scroll.setStyleSheet(
            "QScrollArea { background: transparent; border: none; }"
            "QScrollBar:vertical { width: 6px; background: transparent; }"
            f"QScrollBar::handle:vertical {{ background: {t.scrollbar_handle}; "
            "border-radius: 3px; min-height: 30px; }"
            "QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }"
        )

        content = QWidget()
        content.setStyleSheet("background: transparent;")
        content_layout = QVBoxLayout(content)
        content_layout.setContentsMargins(24, 20, 24, 20)
        content_layout.setSpacing(16)

        self._section_header("📄 选择来源", content_layout)

        fmt_row = QHBoxLayout()
        fmt_row.setSpacing(12)
        self.question_num_fmt = QLineEdit("1.")
        self.question_num_fmt.setPlaceholderText("1.")
        self.question_num_fmt.setFixedHeight(36)
        self.option_prefix = QLineEdit("A.")
        self.option_prefix.setPlaceholderText("A.")
        self.option_prefix.setFixedHeight(36)
        fmt_row.addWidget(self._make_labeled_field("题号格式", self.question_num_fmt))
        fmt_row.addWidget(self._make_labeled_field("选项前缀", self.option_prefix))
        content_layout.addLayout(fmt_row)

        self.word_drop_zone = DropZone("点击或拖拽 .docx 文件", "Word 文档 (*.docx)", theme=t)
        self.word_drop_zone.file_selected.connect(self._on_word_file)
        content_layout.addWidget(self.word_drop_zone)

        self._divider = QFrame()
        self._divider.setFixedHeight(1)
        self._divider.setStyleSheet(f"background: {t.border}; border: none;")
        content_layout.addWidget(self._divider)

        self._section_header("🎨 幻灯片风格", content_layout)

        font_row = QHBoxLayout()
        font_row.setSpacing(12)
        system_fonts = _get_system_fonts()
        self.font_name = QComboBox()
        self.font_name.setEditable(True)
        self.font_name.addItems(system_fonts)
        self.font_name.setCurrentText(system_fonts[0] if system_fonts else "Arial")
        self.font_name.setMinimumWidth(160)
        self.font_name.setFixedHeight(36)
        self.font_name.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        self.font_size = QSpinBox()
        self.font_size.setRange(9, 72)
        self.font_size.setValue(18)
        self.font_size.setMinimumWidth(70)
        self.font_size.setFixedHeight(36)
        self.font_size.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        font_row.addWidget(self._make_labeled_field("字体", self.font_name))
        font_row.addWidget(self._make_labeled_field("字号", self.font_size))
        content_layout.addLayout(font_row)

        spacing_row = QHBoxLayout()
        spacing_row.setSpacing(12)
        self.line_spacing_type = QComboBox()
        self.line_spacing_type.addItems(["1 倍", "1.5 倍", "自定义"])
        self.line_spacing_type.setCurrentText("1 倍")
        self.line_spacing_type.setMinimumWidth(160)
        self.line_spacing_type.setFixedHeight(36)
        self.line_spacing_type.currentTextChanged.connect(self._on_spacing_changed)
        self.line_spacing_value = QDoubleSpinBox()
        self.line_spacing_value.setRange(0.5, 5.0)
        self.line_spacing_value.setSingleStep(0.1)
        self.line_spacing_value.setDecimals(1)
        self.line_spacing_value.setValue(2.0)
        self.line_spacing_value.setFixedHeight(36)
        self.line_spacing_value.setVisible(False)
        self.first_line_indent = QCheckBox("是")
        self.first_line_indent.setChecked(True)
        self.first_line_indent.setFixedHeight(36)
        spacing_row.addWidget(self._make_labeled_field("行间距", self.line_spacing_type))
        spacing_row.addWidget(self._make_labeled_field("自定义", self.line_spacing_value))
        spacing_row.addWidget(self._make_labeled_field("首行缩进", self.first_line_indent))
        spacing_row.addStretch()
        content_layout.addLayout(spacing_row)

        self.ppt_drop_zone = DropZone("点击或拖拽 .pptx 模板", "PPT 模板 (*.pptx)", theme=t)
        self.ppt_drop_zone.file_selected.connect(self._on_ppt_file)
        content_layout.addWidget(self.ppt_drop_zone)

        save_row = QHBoxLayout()
        save_row.setSpacing(12)
        self.out_path_label = QLabel()
        self.out_path_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.out_path_label.setMinimumWidth(200)
        self.out_path_label.setWordWrap(False)
        self._out_path = ""
        self._set_out_path("output.pptx")
        sl = QLabel("保存到:")
        self._save_to_label = sl
        save_row.addWidget(sl)
        save_row.addWidget(self.out_path_label)
        change_btn = AppButton("更改", default_height=28, theme=self.theme)
        change_btn.setFixedWidth(80)
        self._change_btn = change_btn
        change_btn.clicked.connect(lambda: self._on_browse_save(self.out_path_label))
        save_row.addWidget(change_btn)
        content_layout.addLayout(save_row)

        content_layout.addStretch()
        self._scroll.setWidget(content)
        card_layout.addWidget(self._scroll)

        self._root.addWidget(card)

        self.error_label = QLabel("")
        self.error_label.setWordWrap(True)
        self._root.addWidget(self.error_label)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(0)
        self.spinner = QLabel()
        self.spinner.setFixedSize(24, 24)
        self.spinner.setVisible(False)
        self.spinner.setStyleSheet("background: transparent; border: none;")
        btn_row.addWidget(self.spinner)
        self.convert_btn = AnimatedButton("开始转换", default_height=50, theme=self.theme)
        self.convert_btn.clicked.connect(self.on_convert)
        btn_row.addWidget(self.convert_btn)
        self._root.addLayout(btn_row)

        self.progress_bar = AnimatedProgressBar()
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setVisible(False)
        self._root.addWidget(self.progress_bar)

        self.progress_label = QLabel("")
        self.progress_label.setVisible(False)
        self.progress_label.setTextFormat(Qt.RichText)
        self.progress_label.linkActivated.connect(self._open_output_folder)
        self._root.addWidget(self.progress_label)

        self._restyle_all()

    def _section_header(self, text, parent_layout):
        """在指定布局中添加一个章节标题标签。

        Args:
            text: 标题文本。
            parent_layout: 要添加到的布局。
        """
        label = QLabel(text)
        label.setStyleSheet(
            f"font-size: 15px; font-weight: bold; color: {self.theme.card_header_color}; "
            "background: transparent; padding: 0;"
        )
        self._section_labels.append(label)
        parent_layout.addWidget(label)

    def _make_labeled_field(self, label_text, widget):
        """创建一个带标签的表单字段包装器（标签在上，控件在下）。

        Args:
            label_text: 字段标签文本。
            widget: 实际的输入控件。

        Returns:
            包含标签和控件的 QWidget 包装器。
        """
        wrapper = QWidget()
        wrapper.setStyleSheet("background: transparent;")
        layout = QVBoxLayout(wrapper)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(2)
        label = QLabel(label_text)
        self._field_labels.append(label)
        layout.addWidget(label)
        layout.addWidget(widget)
        return wrapper

    # ── actions ─────────────────────────────────────────────────

    def _on_word_file(self, path):
        """Word 文件选择回调，记录路径并将输出文件默认设为同目录下的 output.pptx。"""
        self._word_path = path
        self._update_convert_btn_status()
        out_dir = os.path.dirname(path)
        out_path = os.path.join(out_dir, "output.pptx")
        self._set_out_path(out_path)

    def _on_ppt_file(self, path):
        """PPT 模板文件选择回调，记录路径。"""
        self._ppt_path = path
        self._update_convert_btn_status()

    def _on_spacing_changed(self, text):
        """行间距类型改变时控制自定义数值输入框的可见性。"""
        self.line_spacing_value.setVisible(text == "自定义")

    def _update_convert_btn_status(self):
        """根据当前输入状态更新转换按钮是否可用。"""
        if not getattr(self, '_word_path', None):
            self.convert_btn.set_actionable(False, "请先选择 Word 文档")
            return
        if not getattr(self, '_ppt_path', None):
            self.convert_btn.set_actionable(False, "请先选择 PPT 模板")
            return
        if not self.font_name.currentText().strip():
            self.convert_btn.set_actionable(False, "请选择字体")
            return
        self.convert_btn.set_actionable(True, "")

    def _get_line_spacing_value(self):
        """根据行间距类型返回对应的数值。

        Returns:
            行间距数值（浮点数）。
        """
        return _resolve_line_spacing(
            self.line_spacing_type.currentText(),
            self.line_spacing_value.value(),
        )

    def _on_browse_save(self, label):
        """打开保存文件对话框，让用户选择输出路径。

        Args:
            label: 触发按钮对应的标签（未直接使用，保留兼容性）。
        """
        start_dir = os.path.dirname(self._out_path) if hasattr(self, '_out_path') and self._out_path else ""
        path, _ = QFileDialog.getSaveFileName(
            self, "保存为", start_dir or "output.pptx",
            "PPTX 文件 (*.pptx)"
        )
        if path:
            self._set_out_path(path)

    def _set_out_path(self, path):
        """设置输出文件路径，更新标签文本（省略过长路径）并设置工具提示。

        Args:
            path: 输出文件路径。
        """
        self._out_path = path
        metrics = self.out_path_label.fontMetrics()
        elided = metrics.elidedText(path, Qt.ElideMiddle, 200)
        self.out_path_label.setText(elided)
        self.out_path_label.setToolTip(path)

    def _load_settings(self):
        """从 QSettings 恢复用户上次保存的配置项到各表单控件。"""
        self.question_num_fmt.setText(self.settings.value("question_num_fmt", "1."))
        self.option_prefix.setText(self.settings.value("option_prefix", "A."))
        self.font_name.setCurrentText(self.settings.value("font_name", "微软雅黑"))
        self.font_size.setValue(int(self.settings.value("font_size", 18)))

    def _save_settings(self):
        """将当前表单控件的值持久化到 QSettings。"""
        self.settings.setValue("question_num_fmt", self.question_num_fmt.text())
        self.settings.setValue("option_prefix", self.option_prefix.text())
        self.settings.setValue("font_name", self.font_name.currentText())
        self.settings.setValue("font_size", self.font_size.value())

    def _clear_error(self):
        """清除错误标签的文本。"""
        self.error_label.setText("")

    def _show_error(self, msg):
        """在错误标签中显示指定的错误消息。

        Args:
            msg: 错误消息文本。
        """
        self.error_label.setText(msg)

    def _validate(self):
        """校验用户输入是否完整（Word 文件、PPT 模板、字体是否已选择）。

        Returns:
            True 表示校验通过，False 表示存在错误并已显示。
        """
        errors = []
        if not self._word_path:
            errors.append("请选择 Word 文件")
        if not self._ppt_path:
            errors.append("请选择 PPT 模板")
        if not self.font_name.currentText().strip():
            errors.append("请选择字体")
        if errors:
            self._show_error("；".join(errors))
            return False
        return True

    def on_convert(self):
        """执行转换流程：校验输入 → 后台线程提取题目 → 弹窗预览确认 → 启动生成线程。

        题目提取（解析 Word 文档）在大文档下较重，放到 ExtractWorker 后台线程执行，
        避免阻塞 GUI 造成界面卡顿；提取完成后再弹出预览确认对话框。
        """
        self._clear_error()
        self._save_settings()
        if not self._validate():
            return

        # 进入“识别中”状态：仅显示加载动画，提取在后台线程进行，不阻塞 GUI
        self.convert_btn.set_loading(True)
        self.spinner.setVisible(True)
        self.progress_bar.setVisible(False)
        self.progress_label.setVisible(True)
        self.progress_label.setText("正在识别题目...")

        self._stop_extract_worker()
        self._extract_worker = ExtractWorker(
            self._word_path,
            self.question_num_fmt.text(),
            self.option_prefix.text(),
        )
        self._extract_worker.extracted.connect(self._on_extracted)
        self._extract_worker.error.connect(self._on_extract_error)
        self._extract_worker.start()

    def _reset_extract_ui(self):
        """退出“识别中”状态，复位加载动画与按钮，避免阻塞后续交互。"""
        self.spinner.setVisible(False)
        self.progress_label.setVisible(False)
        self.convert_btn.set_loading(False)

    def _on_extracted(self, questions):
        """后台提取完成回调：复位识别 UI 并弹出预览确认对话框。"""
        self._reset_extract_ui()
        try:
            self._prompt_and_convert(questions)
        finally:
            self._stop_extract_worker()

    def _on_extract_error(self, message):
        """后台提取失败回调：复位 UI 并提示错误。"""
        self._reset_extract_ui()
        self._show_error(f"题目识别失败：{message}")
        logger.error("题目识别失败: %s", message)
        self._stop_extract_worker()

    def _prompt_and_convert(self, questions):
        """展示识别结果预览对话框；用户确认后启动后台生成线程。

        Args:
            questions: 从 Word 文档提取的题目列表（每行一个字符串的二维列表）。
        """
        if not questions:
            self._show_error("未识别到任何题目")
            return

        logger.info("共提取到 %d 道题目", len(questions))

        dlg = QDialog(self)
        dlg.setWindowTitle(f"识别结果 — 共 {len(questions)} 道题")
        dlg.setMinimumSize(620, 450)
        dlg.setModal(True)

        # 构建带 CSS 的 HTML 预览
        font_name = self.font_name.currentText()
        font_size = self.font_size.value()
        line_height = _resolve_line_spacing(
            self.line_spacing_type.currentText(),
            self.line_spacing_value.value(),
        )
        indent_px = round(font_size * 1.333 * 2)
        indent = f"{indent_px}px" if self.first_line_indent.isChecked() else "0"

        css = (
            f"body {{ margin: 24px; font-family: '{_sanitize_font_name(font_name)}'; "
            f"font-size: {font_size}pt; line-height: {line_height}; "
            f"color: {self.theme.text_primary}; "
            f"background-color: {self.theme.card_bg}; }}"
            f".q-header {{ font-weight: bold; color: {self.theme.accent}; "
            f"margin-top: 16px; margin-bottom: 6px; text-indent: 0; }}"
            f".q {{ margin-bottom: 2px; text-indent: {indent}; "
            f"word-break: break-all; }}"
        )

        parts = []
        for i, q in enumerate(questions, 1):
            parts.append(f"<div class='q-header'>第 {i} 题</div>")
            for line in q:
                # 题面文本来自用户文档，必须先转义再渲染，仅保留白名单标签
                parts.append(f"<div class='q'>{_escape_preview_line(line)}</div>")
        html = (
            "<html><head><meta charset='utf-8'>"
            f"<style>{css}</style></head><body>"
            + "\n".join(parts) +
            "</body></html>"
        )

        preview = QTextBrowser()
        preview.setOpenExternalLinks(False)
        preview.setHtml(html)

        # 创建按钮
        btn_box = QDialogButtonBox(
            QDialogButtonBox.Ok | QDialogButtonBox.Cancel,
            Qt.Horizontal,
            dlg,
        )
        btn_box.accepted.connect(dlg.accept)
        btn_box.rejected.connect(dlg.reject)

        # 布局
        dlg_layout = QVBoxLayout(dlg)
        dlg_layout.addWidget(preview)
        dlg_layout.addWidget(btn_box)

        # 显示对话框，等待用户确认
        if dlg.exec() != QDialog.Accepted:
            return

        # 用户确认 → 启动后台生成线程
        out_path = self._out_path

        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, len(questions))
        self.progress_bar.setValue(0)
        self.progress_label.setVisible(True)
        self.progress_label.setText("准备中...")
        self.convert_btn.set_loading(True)
        self.spinner.setVisible(True)

        self._stop_worker()
        self.worker = ConvertWorker(
            self._ppt_path,
            questions,
            font_name,
            font_size,
            out_path,
            self.line_spacing_type.currentText(),
            self._get_line_spacing_value(),
            self.first_line_indent.isChecked(),
        )
        self.worker.progress_text.connect(self._on_progress)
        self.worker.finished.connect(self._on_finished)
        self.worker.start()

    def _stop_worker(self):
        """停止并清理正在运行的后台转换线程，避免重复点击产生孤儿线程。

        CPU 密集线程没有事件循环，quit() 无法中断其 run()，故若仍在运行则
        wait() 等待其自然结束（窗口关闭场景下保证线程被回收，不沦为孤儿），
        随后 deleteLater() 释放对象。先 disconnect 旧信号，避免旧线程的
        残留回调污染 UI 状态。
        """
        w = getattr(self, 'worker', None)
        if w is None:
            return
        try:
            w.progress_text.disconnect()
            w.finished.disconnect()
        except (TypeError, RuntimeError):
            pass
        w.quit()
        if w.isRunning():
            w.wait()
        w.deleteLater()
        self.worker = None

    def _stop_extract_worker(self):
        """停止并清理正在运行的后台题目提取线程（同 _stop_worker 的清理策略）。"""
        w = getattr(self, '_extract_worker', None)
        if w is None:
            return
        try:
            w.extracted.disconnect()
            w.error.disconnect()
        except (TypeError, RuntimeError):
            pass
        w.quit()
        if w.isRunning():
            w.wait()
        w.deleteLater()
        self._extract_worker = None

    def stop_worker(self):
        """供主窗口 closeEvent 调用，统一清理后台转换与题目提取线程。"""
        self._stop_worker()
        self._stop_extract_worker()

    def _on_progress(self, text):
        """处理后台线程的进度信号，更新进度标签和进度条。

        Args:
            text: 进度描述文本。
        """
        self.progress_label.setText(text)
        cur = self.progress_bar.value() + 1
        self.progress_bar.setValueAnimated(cur)

    def _on_finished(self, success, message):
        """处理后台线程的完成信号，隐藏进度 UI 并通过 Toast 显示结果。

        Args:
            success: True 表示转换成功，False 表示失败。
            message: 结果消息或错误信息。
        """
        # 忽略已停止/废弃的旧线程回调，避免重复结果污染 UI 状态
        if self.sender() is not self.worker:
            return
        self.progress_bar.setVisible(False)
        self.progress_label.setVisible(False)
        self.spinner.setVisible(False)
        self.convert_btn.set_loading(False)
        if success:
            self.toast.show_message(message, success=True)
            folder = os.path.dirname(self._out_path) or "."
            self.progress_label.setText(
                f"{message}  <a href=\"folder:{folder}\" style=\"color: inherit; text-decoration: none;\">打开文件夹</a>"
            )
            self.progress_label.setVisible(True)
            logger.info("转换完成: %s", message)
        else:
            self.toast.show_message(message, success=False)
            logger.error("转换失败: %s", message)

        # 释放后台线程对象，避免长期持有造成泄漏
        if self.worker is not None:
            self.worker.deleteLater()
            self.worker = None

    def _open_output_folder(self):
        """使用系统默认文件管理器打开导出文件所在的文件夹。"""
        if not self._out_path:
            return
        folder_path = os.path.dirname(self._out_path)
        try:
            if os.name == 'nt':  # Windows
                subprocess.Popen(['explorer', folder_path])
            elif os.name == 'posix':  # macOS / Linux
                if sys.platform == 'darwin':
                    subprocess.Popen(['open', folder_path])
                else:
                    subprocess.Popen(['xdg-open', folder_path])
        except Exception as e:
            logger.warning("打开文件夹失败: %s", e)

    def on_test(self):
        """创建测试用的 Word 文档和空 PPT 模板，自动填充表单并执行转换流程。

        用于快速验证题目提取和 PPT 生成功能是否正常工作。
        """
        tmp = os.path.join(tempfile.gettempdir(), "quiz2slide_test.docx")
        doc = Document()
        doc.add_paragraph("一些无关的说明文本，不应被识别为题目。")
        doc.add_paragraph("1. 以下哪个是 Python 关键字？")
        doc.add_paragraph("A. class")
        doc.add_paragraph("B. def")
        doc.add_paragraph("C. if")
        doc.add_paragraph("D. all of the above")
        doc.add_paragraph("")
        doc.add_paragraph("2. Python 中列表用什么符号？ A. () B. [] C. {} D. <>")
        doc.add_paragraph("")
        doc.add_paragraph("3. 以下哪些是有效的 JSON 类型？")
        doc.add_paragraph("A. string")
        doc.add_paragraph("B. number")
        doc.add_paragraph("C. boolean")
        doc.add_paragraph("D. null")
        doc.add_paragraph("E. undefined  F. date")
        doc.add_paragraph("")
        doc.add_paragraph("4、以下哪个是动态语言？  A. Python  B. Java  C. C++  D. Rust")
        doc.add_paragraph("")
        doc.add_paragraph("5) Git 中查看提交历史的命令是？")
        doc.add_paragraph("A. git log")
        doc.add_paragraph("B. git status")
        doc.add_paragraph("C. git diff")
        doc.add_paragraph("D. git show")
        doc.save(tmp)
        tpl_tmp = os.path.join(tempfile.gettempdir(), "quiz2slide_test.pptx")
        Presentation().save(tpl_tmp)
        self._word_path = tmp
        self._ppt_path = tpl_tmp
        self._set_out_path(os.path.join(tempfile.gettempdir(), "output.pptx"))
        self.question_num_fmt.setText(r"\d+[.、)]")
        self.option_prefix.setText("A.")
        self.on_convert()

    # ── events ──────────────────────────────────────────────────

    def showEvent(self, event):
        """窗口首次显示时加载用户设置并执行淡入动画。"""
        super().showEvent(event)
        self._load_settings()
        if not hasattr(self, '_faded_in'):
            self._faded_in = True
            self.setWindowOpacity(0.0)
            self._fade_anim = QPropertyAnimation(self, b"windowOpacity")
            self._fade_anim.setDuration(250)
            self._fade_anim.setStartValue(0.0)
            self._fade_anim.setEndValue(1.0)
            self._fade_anim.setEasingCurve(QEasingCurve.OutCubic)
            self._fade_anim.start()



# ── entry ──────────────────────────────────────────────────────────
def get_resource_path(relative_path):
    """获取资源文件路径，兼容开发和 PyInstaller 打包"""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath('.'), relative_path)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon(get_resource_path('icon.png')))
    window = Quiz2SlideTool()
    window.show()

    sys.exit(app.exec())

