"""Similarity Checker 视图 — 仅负责渲染与事件绑定（零业务规则）。

持有 SimilarityViewModel（胶水层），把用户操作翻译成 SimilarityRequest 交给
ViewModel：选择文件 / 设定阈值 → vm.check(request) 后台查重。
进度/结果/失败均通过 vm 信号回流到本视图更新 UI；报告导出为本视图的
展示层职责（把结果格式化写入 .docx）。
"""
from __future__ import annotations

import logging
import os

from docx import Document
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QDoubleSpinBox, QLayout,
    QLineEdit, QFrame, QTextBrowser, QButtonGroup, QRadioButton,
    QFileDialog, QApplication, QPushButton, QSizePolicy, QScrollArea,
)
from PySide6.QtCore import Qt, QSettings
from PySide6.QtGui import QPalette

from widgets import AppButton, AnimatedButton, AnimatedProgressBar, ToastNotification, DropZone, MultiDropZone, StepperInput
from shared.contracts import (
    ManyToManyResult, OneToManyResult, SimilarityMode, SimilarityRequest,
    SimilarityResult,
)
from ui.viewmodels.similarity_viewmodel import SimilarityViewModel
from ui.views.base_view import BaseView
from ui.infra.open_folder import open_folder
from ui.infra.settings_keys import SimilarityKeys
from ui.infra.safe_settings import read_float, read_str

logger = logging.getLogger(__name__)


class SimilarityView(BaseView):
    """题目查重视图。"""

    def get_name(self) -> str:
        return "Similarity Checker"

    def get_nav_title(self) -> str:
        return "🔍 试题查重"

    def get_description(self) -> str:
        return "检测主文档与多个副文档之间的题目重复率，支持精确/模糊匹配，导出查重报告。"

    def __init__(self, view_model: SimilarityViewModel) -> None:
        super().__init__()
        self._vm = view_model
        self._main_path = ""
        self._secondary_paths: list = []
        self._all_paths: list = []
        self._mode = SimilarityMode.ONE_TO_MANY
        self._last_result: SimilarityResult | None = None

        self.settings = QSettings("SimilarityChecker", "SimilarityChecker")

        self._setup_ui()
        self._connect_view_model()
        self._load_settings()
        self.theme.theme_changed.connect(self._on_theme_changed)

    # ── ViewModel 信号绑定（单向数据流：core → UI） ──
    def _connect_view_model(self) -> None:
        self._vm.started.connect(self._on_started)
        self._vm.progress.connect(self._on_progress)
        self._vm.completed.connect(self._on_completed)
        self._vm.failed.connect(self._on_failed)

    # ── 命令转发（单向数据流：UI → core） ──
    def _on_check(self) -> None:
        self._log_browser.clear()
        self._show_result_placeholder()
        self._check_btn.set_loading(True)
        self._export_btn.set_actionable(False, "请先完成检测后导出")
        self._progress_bar.setVisible(True)
        self._progress_bar.setRange(0, 0)
        self._progress_bar.setValue(0)

        if self._mode == SimilarityMode.ONE_TO_MANY:
            if not self._main_path:
                self._log_browser.setHtml(
                    f"<p style='color:{self.theme.error_color};'>请先选择主文档</p>"
                )
                self._finish_check_ui()
                return
            if not self._secondary_paths:
                self._log_browser.setHtml(
                    f"<p style='color:{self.theme.error_color};'>请先选择至少一个副文档</p>"
                )
                self._finish_check_ui()
                return
            req = SimilarityRequest(
                mode=SimilarityMode.ONE_TO_MANY,
                main_path=self._main_path,
                secondary_paths=list(self._secondary_paths),
                threshold=self._threshold_spin.value(),
                num_pattern=self._num_edit.text().strip() or "1.",
                opt_prefix=self._opt_edit.text().strip() or "A.",
            )
        else:
            if len(self._all_paths) < 2:
                self._log_browser.setHtml(
                    f"<p style='color:{self.theme.error_color};'>请先选择至少 2 份文档</p>"
                )
                self._finish_check_ui()
                return
            # mypy 2.3 原生 pydantic 支持把 Field("默认值") 误判为必填（main_path 等），
            # 运行时 pydantic 会正常填充默认值，故忽略该误报。
            req = SimilarityRequest(  # type: ignore[call-arg]
                mode=SimilarityMode.MANY_TO_MANY,
                all_paths=list(self._all_paths),
                threshold=self._threshold_spin.value(),
                num_pattern=self._num_edit.text().strip() or "1.",
                opt_prefix=self._opt_edit.text().strip() or "A.",
            )
        self._vm.check(req)  # 后台执行，结果经 vm 信号回流

    # ── ViewModel 回调 ──
    def _on_started(self, mode: object) -> None:
        self._log_browser.clear()
        self._check_btn.set_loading(True)

    def _on_progress(self, message: str, current: int, total: int) -> None:
        if total > 0:
            self._progress_bar.setRange(0, total)
            self._progress_bar.setValueAnimated(current)
        self._log_browser.append(message)

    def _on_completed(self, result: SimilarityResult) -> None:
        self._finish_check_ui()
        self._last_result = result
        self._render_result(result)
        # 导出按钮严格绑定“检测结果是否有效”，而非“是否检出重复题”。
        # 一次干净的查重（0 重复）同样产生有效报告，应允许导出。
        self._export_btn.set_actionable(self._last_result is not None, "")

    def _on_failed(self, message: object) -> None:
        self._finish_check_ui()
        msg = message if isinstance(message, str) else str(message)
        self._log_browser.append(f"\n错误：{msg}")

    def _finish_check_ui(self) -> None:
        self._check_btn.set_loading(False)
        self._update_check_state()
        self._progress_bar.setVisible(False)

    # ── 结果展示（结构化卡片，高亮核心数据） ──
    def _render_result(self, result: SimilarityResult) -> None:
        """以结构化卡片渲染检测结果，高亮重复率 / 重复题目数。"""
        self._clear_layout(self._result_layout)
        if isinstance(result, ManyToManyResult):
            self._build_many_to_many_cards(result)
        else:
            self._build_one_to_many_cards(result)

    def _clear_layout(self, layout: QLayout) -> None:
        while layout.count():
            item = layout.takeAt(0)
            if item is None:
                continue
            w = item.widget()
            if w is not None:
                w.setParent(None)
                w.deleteLater()
            else:
                sub = item.layout()
                if sub is not None:
                    self._clear_layout(sub)

    def _show_result_placeholder(self) -> None:
        self._clear_layout(self._result_layout)
        self._result_placeholder = QLabel(
            "完成检测后，将在此以结构化卡片展示重复率与重复题目。"
        )
        self._result_placeholder.setWordWrap(True)
        self._result_layout.addWidget(self._result_placeholder)

    def _build_one_to_many_cards(self, result: OneToManyResult) -> None:
        t = self.theme
        total = result.main_count
        dup = result.duplicate_count
        rate = dup / max(total, 1) * 100
        rate_color = t.danger if rate >= 50 else t.accent

        summary = self._card()
        sl = QVBoxLayout(summary)
        sl.setContentsMargins(self.theme.page_pad_y, self.theme.spacing, self.theme.page_pad_y, self.theme.spacing)
        sl.setSpacing(self.theme.control_spacing)
        title = QLabel("检测摘要")
        title.setStyleSheet(t.qss_section_header())
        rate_lbl = QLabel(f"{rate:.1f}%")
        rate_lbl.setStyleSheet(
            f"font-size: 28px; font-weight: bold; color: {rate_color};"
        )
        sub_lbl = QLabel(f"重复题目 {dup} 道 / 总题目 {total} 道")
        sub_lbl.setStyleSheet(f"font-size: 12px; color: {t.text_secondary};")
        sl.addWidget(title)
        sl.addWidget(rate_lbl)
        sl.addWidget(sub_lbl)
        self._result_layout.addWidget(summary)

        if not result.details:
            empty = QLabel("未检测到重复题目，题库质量良好。")
            empty.setWordWrap(True)
            empty.setStyleSheet(f"font-size: 12px; color: {t.text_secondary};")
            self._result_layout.addWidget(empty)
            return

        for d in result.details:
            dc = self._card()
            dl = QVBoxLayout(dc)
            dl.setContentsMargins(16, 12, 16, 12)
            dl.setSpacing(6)
            h = QLabel(f"第 {d.index} 题")
            h.setStyleSheet(
                f"font-size: 13px; font-weight: bold; color: {t.text_primary};"
            )
            dl.addWidget(h)
            text0 = d.text[0] if d.text else ""
            preview = (text0[:50] + "…") if len(text0) > 50 else text0
            pv = QLabel(preview)
            pv.setWordWrap(True)
            pv.setStyleSheet(f"font-size: 12px; color: {t.text_secondary};")
            dl.addWidget(pv)
            for item in d.sources:
                s = QLabel(
                    f"重复来源：{item.file}（相似度 {item.score:.2f}，{item.reason}）"
                )
                s.setWordWrap(True)
                s.setStyleSheet(
                    f"font-size: 12px; color: {t.text_secondary}; padding-left: 12px;"
                )
                dl.addWidget(s)
            self._result_layout.addWidget(dc)

    def _build_many_to_many_cards(self, result: ManyToManyResult) -> None:
        t = self.theme
        internal = sum(1 for p in result.duplicate_pairs if p.pair_type == "internal")
        cross = sum(1 for p in result.duplicate_pairs if p.pair_type == "cross")
        total_pairs = len(result.duplicate_pairs)

        summary = self._card()
        sl = QVBoxLayout(summary)
        sl.setContentsMargins(self.theme.page_pad_y, self.theme.spacing, self.theme.page_pad_y, self.theme.spacing)
        sl.setSpacing(self.theme.control_spacing)
        title = QLabel("检测摘要")
        title.setStyleSheet(t.qss_section_header())
        rate_lbl = QLabel(f"{total_pairs} 对")
        rate_lbl.setStyleSheet(
            f"font-size: 28px; font-weight: bold; color: {t.accent};"
        )
        sub_lbl = QLabel(
            f"文档数 {result.document_count} · 总题目 {result.total_questions} · "
            f"重复对 {total_pairs}（跨文档 {cross} / 文档内 {internal}）"
        )
        sub_lbl.setWordWrap(True)
        sub_lbl.setStyleSheet(f"font-size: 12px; color: {t.text_secondary};")
        sl.addWidget(title)
        sl.addWidget(rate_lbl)
        sl.addWidget(sub_lbl)
        self._result_layout.addWidget(summary)

        if result.doc_questions:
            dist = self._card()
            dl = QVBoxLayout(dist)
            dl.setContentsMargins(16, 12, 16, 12)
            dl.setSpacing(4)
            dh = QLabel("文档题目分布")
            dh.setStyleSheet(
                f"font-size: 13px; font-weight: bold; color: {t.text_primary};"
            )
            dl.addWidget(dh)
            for fname, count in result.doc_questions.items():
                row = QLabel(f"{fname}：{count} 题")
                row.setStyleSheet(f"font-size: 12px; color: {t.text_secondary};")
                dl.addWidget(row)
            self._result_layout.addWidget(dist)

        for i, pair in enumerate(result.duplicate_pairs, 1):
            tag = "跨文档" if pair.pair_type == "cross" else "文档内"
            pc = self._card()
            pl = QVBoxLayout(pc)
            pl.setContentsMargins(16, 12, 16, 12)
            pl.setSpacing(6)
            h = QLabel(
                f"{i}. [{tag}] {pair.q1_file}-第{pair.q1_index}题 "
                f"⇄ {pair.q2_file}-第{pair.q2_index}题 "
                f"（{pair.score:.2f}，{pair.reason}）"
            )
            h.setWordWrap(True)
            h.setStyleSheet(
                f"font-size: 13px; font-weight: bold; color: {t.text_primary};"
            )
            pl.addWidget(h)
            q1 = pair.q1_text[0] if pair.q1_text else ""
            preview = (q1[:40] + "…") if len(q1) > 40 else q1
            pv = QLabel(preview)
            pv.setWordWrap(True)
            pv.setStyleSheet(f"font-size: 12px; color: {t.text_secondary};")
            pl.addWidget(pv)
            self._result_layout.addWidget(pc)

    def _card(self) -> QFrame:
        card = QFrame()
        card.setStyleSheet(self.theme.qss_card())
        return card

    # ── 报告导出（展示层职责） ──
    def _on_export(self) -> None:
        if self._last_result is None:
            return
        start_dir = os.path.dirname(self._main_path or (self._all_paths[0] if self._all_paths else ""))
        default_name = os.path.join(start_dir, "查重报告.docx")
        path, _ = QFileDialog.getSaveFileName(
            self, "导出查重报告", default_name, "Word 文档 (*.docx)"
        )
        if not path:
            return
        if isinstance(self._last_result, ManyToManyResult):
            self._export_many_to_many(self._last_result, path)
        else:
            self._export_one_to_many(self._last_result, path)
        self._open_folder(path)
        folder = os.path.dirname(path) or "."
        self._export_status.setText(
            f"✅ 报告已导出：{os.path.basename(path)}　"
            f"<a href=\"folder:{folder}\" style=\"color: {self.theme.accent}; "
            f"text-decoration: underline;\">打开文件夹</a>"
        )
        self._export_status.setVisible(True)
        self.toast.show_message("查重报告已导出", success=True)

    def _export_one_to_many(self, result: OneToManyResult, path: str) -> None:
        doc = Document()
        doc.add_heading("题目查重报告（1对多模式）", 0)
        doc.add_paragraph(
            f"主文档题目数：{result.main_count}，"
            f"重复题目数：{result.duplicate_count}，"
            f"重复率：{result.duplicate_count / max(result.main_count, 1) * 100:.1f}%"
        )
        for d in result.details:
            doc.add_heading(f"第 {d.index} 题", 2)
            for line in d.text:
                doc.add_paragraph(line, style="List Bullet")
            doc.add_paragraph(
                "重复来源：" + "; ".join(
                    f"{item.file} ({item.score:.2f}, {item.reason})"
                    for item in d.sources
                )
            )
        doc.save(path)

    def _export_many_to_many(self, result: ManyToManyResult, path: str) -> None:
        doc = Document()
        internal = sum(1 for p in result.duplicate_pairs if p.pair_type == "internal")
        cross = sum(1 for p in result.duplicate_pairs if p.pair_type == "cross")
        doc.add_heading("题目查重报告（多对多模式）", 0)
        doc.add_paragraph(
            f"文档数：{result.document_count}，"
            f"总题目数：{result.total_questions}，"
            f"重复对总数：{len(result.duplicate_pairs)}"
            f"（文档内 {internal}，跨文档 {cross}）"
        )
        doc.add_heading("文档题目分布", 2)
        for fname, count in result.doc_questions.items():
            doc.add_paragraph(f"{fname}：{count} 题", style="List Bullet")
        doc.add_heading("重复详情", 2)
        for i, pair in enumerate(result.duplicate_pairs, 1):
            tag = "跨文档" if pair.pair_type == "cross" else "文档内"
            doc.add_heading(
                f"{i}. [{tag}] {pair.q1_file}-第{pair.q1_index}题 "
                f"⇄ {pair.q2_file}-第{pair.q2_index}题 "
                f"({pair.score:.2f}, {pair.reason})",
                3,
            )
            doc.add_paragraph(f"题目 1（{pair.q1_file} 第{pair.q1_index}题）：")
            for line in pair.q1_text:
                doc.add_paragraph(line, style="List Bullet")
            doc.add_paragraph(f"题目 2（{pair.q2_file} 第{pair.q2_index}题）：")
            for line in pair.q2_text:
                doc.add_paragraph(line, style="List Bullet")
            doc.add_paragraph("")
        doc.save(path)

    # ── UI 构建 ──
    def _setup_ui(self) -> None:
        t = self.theme
        root = QVBoxLayout(self)
        root.setContentsMargins(self.theme.page_pad_x, self.theme.page_pad_y, self.theme.page_pad_x, self.theme.page_pad_y)
        root.setSpacing(0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self._scroll = scroll

        content = QWidget()
        content.setStyleSheet("background: transparent;")
        content_layout = QVBoxLayout(content)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(self.theme.spacing)

        # 查重模式
        mode_row = QHBoxLayout()
        mode_row.setSpacing(self.theme.control_spacing)
        mode_label = QLabel("查重模式：")
        mode_label.setStyleSheet(
            f"font-size: 13px; font-weight: bold; color: {t.card_header_color}; "
            "background: transparent; padding: 0;"
        )
        mode_row.addWidget(mode_label)
        self._mode_group = QButtonGroup(self)
        self._radio_1toN = QRadioButton("1对多")
        self._radio_NtoN = QRadioButton("多对多")
        self._radio_1toN.setChecked(True)
        self._mode_group.addButton(self._radio_1toN, 0)
        self._mode_group.addButton(self._radio_NtoN, 1)
        self._mode_group.buttonClicked.connect(self._on_mode_changed)
        mode_row.addWidget(self._radio_1toN)
        mode_row.addWidget(self._radio_NtoN)
        mode_row.addStretch()
        content_layout.addLayout(mode_row)

        # 主文档（模块卡片，主色边框强调）
        self._one_to_many_widgets = []
        card_main, lm = self._make_module_card("主文档（选择题库）")
        self._main_drop_zone = DropZone(
            "点击或拖拽 .docx 文件", "Word 文档 (*.docx)", theme=t, variant="primary"
        )
        self._main_drop_zone.file_selected.connect(self._on_main_file)
        self._main_drop_zone.file_cleared.connect(self._on_main_cleared)
        self._main_drop_zone.invalid_file.connect(self._on_invalid_file)
        lm.addWidget(self._main_drop_zone)
        self._one_to_many_widgets.append(card_main)
        content_layout.addWidget(card_main)

        self._divider = QFrame()
        self._divider.setFixedHeight(1)
        self._one_to_many_widgets.append(self._divider)
        content_layout.addWidget(self._divider)

        # 副文档（模块卡片，中性边框）
        card_sec, ls = self._make_module_card("副文档（对比库，可多选）")
        self._secondary_drop_zone = MultiDropZone(
            "点击或拖拽 .docx 文件（可多选）", "Word 文档 (*.docx)",
            theme=t, variant="secondary"
        )
        self._secondary_drop_zone.files_selected.connect(self._on_secondary_files)
        self._secondary_drop_zone.invalid_file.connect(self._on_invalid_file)
        ls.addWidget(self._secondary_drop_zone)
        self._one_to_many_widgets.append(card_sec)
        content_layout.addWidget(card_sec)

        # 所有文档（多对多，中性边框）
        self._many_to_many_widgets = []
        card_all, la = self._make_module_card("所有文档（可多选，至少 2 份）")
        self._all_drop_zone = MultiDropZone(
            "点击或拖拽 .docx 文件（可多选）", "Word 文档 (*.docx)",
            theme=t, variant="secondary"
        )
        self._all_drop_zone.files_selected.connect(self._on_all_files)
        self._all_drop_zone.invalid_file.connect(self._on_invalid_file)
        la.addWidget(self._all_drop_zone)
        self._many_to_many_widgets.append(card_all)
        content_layout.addWidget(card_all)
        for w in self._many_to_many_widgets:
            w.setVisible(False)

        # 参数网格（阈值 / 题号格式 / 选项前缀 + 重置）
        params_row = QHBoxLayout()
        params_row.setSpacing(12)
        self._threshold_spin = QDoubleSpinBox()
        self._threshold_spin.setRange(0.5, 1.0)
        self._threshold_spin.setSingleStep(0.01)
        self._threshold_spin.setDecimals(2)
        self._threshold_spin.setValue(read_float(self.settings, SimilarityKeys.THRESHOLD, 0.8, lo=0.5, hi=1.0))
        self._threshold_spin.setFixedHeight(36)
        self._num_edit = QLineEdit()
        self._num_edit.setPlaceholderText("如 1.")
        self._num_edit.setFixedHeight(36)
        self._opt_edit = QLineEdit()
        self._opt_edit.setPlaceholderText("如 A.")
        self._opt_edit.setFixedHeight(36)

        # 相似度阈值步进控件：左侧− / 中间输入 / 右侧+，与 Quiz2Slide 字号/行距 1:1 统一；
        # 内部 QDoubleSpinBox 原生上下箭头已在 StepperInput 中彻底移除。
        self._threshold_stepper = StepperInput(
            spin=self._threshold_spin, theme=t, minus_text="−", plus_text="+"
        )

        params_row.addWidget(self._make_labeled_field("相似度阈值", self._threshold_stepper))
        params_row.addWidget(self._make_labeled_field("题号格式", self._num_edit))
        params_row.addWidget(self._make_labeled_field("选项前缀", self._opt_edit))
        params_row.addStretch()
        self._reset_btn = AppButton(
            "重置", default_height=32, theme=self.theme, variant="secondary"
        )
        self._reset_btn.clicked.connect(self._on_reset_settings)
        params_row.addWidget(self._reset_btn, alignment=Qt.AlignmentFlag.AlignVCenter)
        content_layout.addLayout(params_row)

        # 连 StepperInput 对外信号而非内部 QDoubleSpinBox：
        # 既覆盖 ± 按钮与键盘输入两条路径，也不再依赖控件内部实现（消除封装泄漏）。
        # 此前此处额外给 ± 按钮接了一份 stepUp/stepDown，与 StepperInput 内置接线
        # 叠加导致每次点击步进两格（0.8 → 0.82），已随之修复。
        self._threshold_stepper.valueChanged.connect(self._save_settings)
        self._num_edit.textChanged.connect(self._save_settings)
        self._opt_edit.textChanged.connect(self._save_settings)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(self.theme.control_spacing)
        self._check_btn = AnimatedButton(
            "开始检测", default_height=40, theme=self.theme, loading_text="检测中..."
        )
        self._check_btn.clicked.connect(self._on_check)
        self._export_btn = AppButton(
            "导出报告", default_height=32, theme=self.theme, variant="secondary"
        )
        self._export_btn.clicked.connect(self._on_export)
        self._export_btn.set_actionable(False, "请先完成检测后导出")
        btn_row.addWidget(self._check_btn)
        btn_row.addWidget(self._export_btn)
        content_layout.addLayout(btn_row)

        content_layout.addSpacing(self.theme.control_spacing)

        self._progress_bar = AnimatedProgressBar()
        self._progress_bar.setTextVisible(False)
        self._progress_bar.setVisible(False)
        content_layout.addWidget(self._progress_bar)

        # 处理日志（检测过程中的步骤反馈）
        self._log_browser = QTextBrowser()
        self._log_browser.setOpenExternalLinks(False)
        self._log_browser.setMinimumHeight(90)
        self._log_browser.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        content_layout.addWidget(self._log_browser)

        # 检测结果（结构化卡片，替代纯文本展示）
        self._result_area = QFrame()
        self._result_area.setObjectName("result_area")
        self._result_layout = QVBoxLayout(self._result_area)
        self._result_layout.setContentsMargins(0, 0, 0, 0)
        self._result_layout.setSpacing(12)
        self._show_result_placeholder()
        content_layout.addWidget(self._result_area)

        # 导出后状态：可点击「打开文件夹」链接
        self._export_status = QLabel("")
        self._export_status.setVisible(False)
        self._export_status.setTextFormat(Qt.TextFormat.RichText)
        self._export_status.linkActivated.connect(self._open_folder_link)
        content_layout.addWidget(self._export_status)

        scroll.setWidget(content)
        root.addWidget(scroll, 1)
        self.toast = ToastNotification(self, theme=self.theme)
        self._update_check_state()
        self._restyle_all()

    def _on_mode_changed(self, button: object) -> None:
        if button == self._radio_1toN:
            self._mode = SimilarityMode.ONE_TO_MANY
            for w in self._one_to_many_widgets:
                w.setVisible(True)
            for w in self._many_to_many_widgets:
                w.setVisible(False)
        else:
            self._mode = SimilarityMode.MANY_TO_MANY
            for w in self._one_to_many_widgets:
                w.setVisible(False)
            for w in self._many_to_many_widgets:
                w.setVisible(True)
        self._log_browser.clear()
        self._update_check_state()

    def _on_main_file(self, path: str) -> None:
        self._main_path = path
        self._log_browser.clear()
        self._update_check_state()

    def _on_main_cleared(self) -> None:
        self._main_path = ""
        self._update_check_state()

    def _on_secondary_files(self, paths: list) -> None:
        self._secondary_paths = paths
        self._log_browser.clear()
        self._update_check_state()

    def _on_all_files(self, paths: list) -> None:
        self._all_paths = paths
        self._log_browser.clear()
        self._update_check_state()

    def _on_invalid_file(self, path: str) -> None:
        self.toast.show_message(f"文件格式不支持：{os.path.basename(path)}", success=False)

    def _update_check_state(self) -> None:
        """依据模式与已选文件，启用/置灰「开始检测」按钮。"""
        if self._check_btn._loading:
            return
        if self._mode == SimilarityMode.ONE_TO_MANY:
            if not self._main_path and not self._secondary_paths:
                self._check_btn.set_actionable(False, "请先选择主文档与至少一个副文档")
            elif not self._main_path:
                self._check_btn.set_actionable(False, "请先选择主文档")
            elif not self._secondary_paths:
                self._check_btn.set_actionable(False, "请先选择至少一个副文档")
            else:
                self._check_btn.set_actionable(True, "")
        else:
            if len(self._all_paths) < 2:
                self._check_btn.set_actionable(False, "请先选择至少 2 份文档")
            else:
                self._check_btn.set_actionable(True, "")

    def _open_folder_link(self, link: str) -> None:
        if link.startswith("folder:"):
            self._open_folder(link[7:])

    # ── QSettings 持久化（P1 #4） ──
    def _load_settings(self) -> None:
        # 加载期间屏蔽 textChanged/valueChanged，避免部分字段尚未载入时就触发
        # _save_settings 把“半载状态”（如仍为默认的 opt_edit）写回，覆盖已存值。
        self._num_edit.blockSignals(True)
        self._opt_edit.blockSignals(True)
        self._threshold_spin.blockSignals(True)
        self._num_edit.setText(read_str(self.settings, SimilarityKeys.NUM_PATTERN, "1."))
        self._opt_edit.setText(read_str(self.settings, SimilarityKeys.OPT_PREFIX, "A."))
        self._threshold_spin.setValue(read_float(self.settings, SimilarityKeys.THRESHOLD, 0.8, lo=0.5, hi=1.0))
        self._num_edit.blockSignals(False)
        self._opt_edit.blockSignals(False)
        self._threshold_spin.blockSignals(False)

    def _save_settings(self) -> None:
        self.settings.setValue(SimilarityKeys.THRESHOLD, self._threshold_spin.value())
        self.settings.setValue(SimilarityKeys.NUM_PATTERN, self._num_edit.text())
        self.settings.setValue(SimilarityKeys.OPT_PREFIX, self._opt_edit.text())

    def _on_reset_settings(self) -> None:
        self._threshold_spin.setValue(0.8)
        self._num_edit.setText("1.")
        self._opt_edit.setText("A.")
        self._save_settings()

    def _on_theme_changed(self) -> None:
        self._restyle_all()

    def _open_folder(self, path: str) -> None:
        folder = os.path.dirname(path) or "."
        open_folder(folder)

    def _restyle_all(self) -> None:
        t = self.theme
        pal = self.palette()
        pal.setColor(QPalette.ColorRole.Window, t.window_solid_bg)
        self.setPalette(pal)
        self.setAutoFillBackground(True)

        for card in self._module_cards:
            card.setStyleSheet(t.qss_card())
        self._result_area.setStyleSheet("background: transparent; border: none;")

        for dz in (self._main_drop_zone, self._secondary_drop_zone, self._all_drop_zone):
            dz._theme = t
            dz._apply_style()

        radio_style = (
            f"QRadioButton {{ color: {t.text_primary}; font-size: 13px; "
            f"background: transparent; spacing: 6px; padding: 4px 12px; }}"
            f"QRadioButton::indicator {{ width: 16px; height: 16px; }}"
        )
        self._radio_1toN.setStyleSheet(radio_style)
        self._radio_NtoN.setStyleSheet(radio_style)

        header_s = t.qss_section_header()
        for lbl in self._section_labels:
            lbl.setStyleSheet(header_s)

        self._divider.setStyleSheet(t.qss_divider())

        label_s = f"font-size: 12px; color: {t.text_secondary}; margin-bottom: 2px;"
        for lbl in self._field_labels:
            lbl.setStyleSheet(label_s)

        self._check_btn.set_theme(t)
        self._export_btn.set_theme(t)
        self._reset_btn.set_theme(t)
        self._progress_bar.setStyleSheet(t.qss_progress_bar())

        self._log_browser.setStyleSheet(
            f"QTextBrowser {{ background: {t.input_bg}; color: {t.text_primary}; "
            f"border: none; border-radius: {t.radius}px; padding: 12px; font-size: 13px; }}"
        )
        # 文本框（题号格式 / 选项前缀）：全局规范，不再包含 QDoubleSpinBox 部分
        input_s = (
            f"QLineEdit {{ padding: 4px 8px; border: 1px solid transparent; "
            f"border-radius: {t.radius}px; font-size: 13px; background: {t.input_bg}; "
            f"color: {t.text_primary}; }}"
            f"QLineEdit:hover {{ border-color: {t.accent}; }}"
            f"QLineEdit:focus {{ border: 1px solid {t.accent}; background: {t.card_bg}; }}"
        )
        self._num_edit.setStyleSheet(input_s)
        self._opt_edit.setStyleSheet(input_s)

        # 相似度阈值步进控件（中间输入 + 两侧按钮）统一应用全局规范
        self._threshold_stepper.set_theme(t)

        self._export_status.setStyleSheet(f"color: {t.text_secondary}; font-size: 12px;")

        # 结果区：如有结果则按当前主题重渲染，否则刷新占位提示样式
        if self._last_result is not None:
            self._render_result(self._last_result)
        else:
            self._show_result_placeholder()

        self._scroll.setStyleSheet(t.qss_scrollbar())

    def stop_worker(self) -> None:
        """供主窗口 closeEvent 调用，取消正在运行的后台任务。"""
        self._vm.cancel_current()
